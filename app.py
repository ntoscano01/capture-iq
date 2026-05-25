"""
CaptureIQ - Flask Web Application
Local/hosted app for managing SBIR solicitation pipeline and capture.
"""

import os
import threading
from datetime import datetime
import csv
import io
from functools import wraps
from flask import Flask, render_template, request, redirect, url_for, jsonify, flash, Response, send_file, abort
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user

import database as db

# Allow OAuth2 over plain HTTP for local development
os.environ.setdefault("OAUTHLIB_INSECURE_TRANSPORT", "1")

app = Flask(__name__)
app.secret_key = os.environ.get("CAPTUREIQ_SECRET_KEY", "captureiq-local-secret-change-me")

# Session configuration
app.config['SESSION_COOKIE_SECURE'] = False  # Set to True in production (HTTPS only)
app.config['SESSION_COOKIE_HTTPONLY'] = True  # Prevent JavaScript access to session cookie
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'  # CSRF protection

# ── Flask-Login setup ─────────────────────────────────────────────────────────

login_manager = LoginManager(app)
login_manager.login_view = "login"
login_manager.login_message = "Please log in to access CaptureIQ."
login_manager.login_message_category = "warning"


class User(UserMixin):
    def __init__(self, id, username, email, role, is_active=True, is_capture_manager=False):
        self.id = id
        self.username = username
        self.email = email
        self.role = role
        self._active = bool(is_active)
        self.is_capture_manager = bool(is_capture_manager)

    @property
    def is_active(self):
        return self._active

    @property
    def is_admin(self):
        return self.role == "admin"


@login_manager.user_loader
def load_user(user_id):
    row = db.get_user_by_id(int(user_id))
    if row:
        return User(row["id"], row["username"], row["email"],
                    row["role"], row["is_active"], row.get("is_capture_manager", 0))
    return None

# Track background ingestion jobs
_jobs: dict[str, dict] = {}


# ── Jinja2 helpers ─────────────────────────────────────────────────────────────

def _file_icon(filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return {
        "pdf": "bi-file-earmark-pdf",
        "doc": "bi-file-earmark-word", "docx": "bi-file-earmark-word",
        "xls": "bi-file-earmark-excel", "xlsx": "bi-file-earmark-excel",
        "ppt": "bi-file-earmark-ppt", "pptx": "bi-file-earmark-ppt",
        "txt": "bi-file-earmark-text", "md": "bi-file-earmark-text",
        "csv": "bi-file-earmark-spreadsheet",
        "png": "bi-file-earmark-image", "jpg": "bi-file-earmark-image",
        "jpeg": "bi-file-earmark-image", "gif": "bi-file-earmark-image",
        "zip": "bi-file-earmark-zip",
        "msg": "bi-envelope", "eml": "bi-envelope",
    }.get(ext, "bi-file-earmark")


def _activity_icon(event_type: str) -> str:
    return {
        "created":      "bi-plus-circle-fill",
        "stage_change": "bi-arrow-right-circle-fill",
        "checklist":    "bi-check-circle-fill",
        "file_upload":  "bi-upload",
        "file_delete":  "bi-trash3",
        "note":         "bi-chat-left-text-fill",
    }.get(event_type, "bi-dot")


def _activity_color(event_type: str) -> str:
    return {
        "created":      "#0d6efd",
        "stage_change": "#6f42c1",
        "checklist":    "#198754",
        "file_upload":  "#0d6efd",
        "file_delete":  "#dc3545",
        "note":         "#fd7e14",
    }.get(event_type, "#adb5bd")


app.jinja_env.globals.update(
    file_icon=_file_icon,
    activity_icon=_activity_icon,
    activity_color=_activity_color,
)

# ── Access Control Decorator ──────────────────────────────────────────────────────

def require_project_access(required_role='viewer'):
    """
    Decorator to enforce project-level access control.
    required_role can be 'viewer', 'editor', or 'owner'.
    Roles inherit: owner > editor > viewer
    """
    def decorator(f):
        @wraps(f)
        @login_required
        def decorated_function(project_id, *args, **kwargs):
            project = db.get_project(project_id)
            if not project:
                abort(404)

            # Check if current user is owner
            user_role = db.get_project_member_role(project_id, current_user.id)
            if user_role is None and project.get('owner_id') != current_user.id:
                # User has no access
                flash("You do not have access to this project.", "danger")
                abort(403)

            # Determine actual role
            actual_role = user_role if user_role else 'owner'

            # Check role level
            role_levels = {'owner': 3, 'editor': 2, 'viewer': 1}
            if role_levels.get(actual_role, 0) < role_levels.get(required_role, 0):
                flash(f"You need {required_role} access to perform this action.", "danger")
                abort(403)

            # Store role in g for use in the route
            from flask import g
            g.project_user_role = actual_role

            return f(project_id, *args, **kwargs)
        return decorated_function
    return decorator


def require_capture_manager(f):
    """
    Decorator to enforce that only capture managers can access a route.
    Also allows admins to access capture manager routes.
    """
    @wraps(f)
    @login_required
    def decorated_function(*args, **kwargs):
        if not current_user.is_capture_manager and not current_user.is_admin:
            flash("You do not have permission to access capture management features.", "danger")
            abort(403)
        return f(*args, **kwargs)
    return decorated_function


@app.context_processor
def inject_gdrive_status():
    try:
        from integrations import google_drive as gd
        return {
            "gdrive_connected": gd.is_connected(),
            "gdrive_has_creds": gd.has_credentials_file(),
        }
    except Exception:
        return {"gdrive_connected": False, "gdrive_has_creds": False}


@app.context_processor
def inject_user_helpers():
    """Inject helper functions for user role checking in templates."""
    def user_is_capture_manager():
        return current_user.is_capture_manager if current_user.is_authenticated else False

    task_counts = None
    if current_user.is_authenticated and not current_user.is_admin:
        try:
            task_counts = db.get_task_counts_for_user(current_user.id)
        except Exception:
            task_counts = {"active": 0, "overdue": 0}

    return dict(user_is_capture_manager=user_is_capture_manager, task_counts=task_counts)

# File upload configuration
UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), "project_uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ALLOWED_EXTENSIONS = {
    "pdf", "doc", "docx", "xls", "xlsx", "ppt", "pptx",
    "txt", "md", "csv", "png", "jpg", "jpeg", "gif",
    "zip", "msg", "eml",
}

def _allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

# ── Shutdown ───────────────────────────────────────────────────────────────────

@app.route("/shutdown", methods=["POST"])
@login_required
def shutdown():
    """Shut down the local server. Available to all logged-in users."""
    import signal, threading
    threading.Timer(0.5, lambda: os.kill(os.getpid(), signal.SIGTERM)).start()
    return render_template("shutdown.html")


# ── Startup ────────────────────────────────────────────────────────────────────

@app.before_request
def ensure_db():
    """Ensure DB is initialized (runs once on first request)."""
    if not hasattr(app, "_db_ready"):
        db.init_db()
        app._db_ready = True


# Endpoints that admins cannot access (redirect to admin panel)
_USER_ONLY_ENDPOINTS = {
    "dashboard", "topics", "topic_detail", "toggle_topic_favorite",
    "set_topic_score", "set_topic_notes", "set_topic_status",
    "export_topic_pdf", "export_topic_docx", "export_topics_csv",
    "ingest_page", "ingest_sbir_gov", "ingest_sbir_topics",
    "ingest_navy", "ingest_dod", "delete_ingest_log",
    "analytics", "search",
    "projects", "create_project", "project_detail", "edit_project",
    "set_project_stage", "delete_project",
    "upload_project_file", "download_project_file", "delete_project_file",
    "toggle_checklist", "add_checklist_item", "delete_checklist_item",
    # Collaboration endpoints
    "share_project", "remove_project_member", "update_member_role",
    "add_comment", "delete_comment",
    "get_notifications", "mark_notification_read", "mark_all_read", "notifications",
    "link_document", "remove_document",
    "gdrive_settings", "gdrive_connect", "gdrive_callback", "gdrive_disconnect",
    "sharepoint_settings", "sharepoint_disconnect",
    # Task management
    "tasks", "create_task", "edit_task", "delete_task", "update_task_status",
    "update_checklist_schedule",
    # Capture plan actions
    "link_project_from_capture", "add_key_contact", "delete_key_contact",
    "all_key_contacts",
    # Ingest jobs
    "ingest_jobs", "create_ingest_job", "edit_ingest_job", "toggle_ingest_job", "delete_ingest_job",
    "run_ingest_job_now",
}

# Endpoints that only admins can access
_ADMIN_ONLY_ENDPOINTS = {
    "admin_users", "admin_create_user", "admin_toggle_user",
    "admin_delete_user", "admin_reset_password", "admin_unlock_user",
    "admin_audit_log", "admin_audit_log_export",
    "admin_db", "admin_db_backup", "admin_db_import", "admin_db_purge_topics",
    "admin_settings", "admin_save_settings", "admin_save_smtp",
    "admin_pending_registrations", "admin_approve_registration", "admin_reject_registration",
}


@app.before_request
def enforce_role_routing():
    """Redirect admins away from user pages and users away from admin pages."""
    if not current_user.is_authenticated:
        return None
    endpoint = request.endpoint
    if not endpoint:
        return None
    if current_user.is_admin and endpoint in _USER_ONLY_ENDPOINTS:
        return redirect(url_for("admin_users"))
    if not current_user.is_admin and endpoint in _ADMIN_ONLY_ENDPOINTS:
        flash("Admin access required.", "danger")
        return redirect(url_for("dashboard"))


# ── Auth — Login / Logout / Setup ─────────────────────────────────────────────

@app.route("/login", methods=["GET", "POST"])
def login():
    # Redirect to setup if no users exist yet
    if db.count_users() == 0:
        return redirect(url_for("setup"))
    if current_user.is_authenticated:
        dest = url_for("admin_users") if current_user.is_admin else url_for("dashboard")
        return redirect(dest)
    if request.method == "POST":
        username  = request.form.get("username", "").strip()
        password  = request.form.get("password", "")
        ip        = request.remote_addr or "unknown"
        row       = db.get_user_by_username(username)

        if not row:
            db.write_audit_log("LOGIN_FAILED", username=username,
                               detail="Unknown user", ip_address=ip)
            flash("Invalid username or password.", "danger")
            return render_template("login.html")

        if not row["is_active"]:
            db.write_audit_log("LOGIN_FAILED", username=username, user_id=row["id"],
                               detail="Account inactive", ip_address=ip)
            flash("Your account has been deactivated. Contact an admin.", "danger")
            return render_template("login.html")

        if row.get("locked_at"):
            db.write_audit_log("LOGIN_FAILED", username=username, user_id=row["id"],
                               detail="Account locked", ip_address=ip)
            flash("Account locked due to too many failed attempts. Contact an admin.", "danger")
            return render_template("login.html")

        if check_password_hash(row["password_hash"], password):
            user = User(row["id"], row["username"], row["email"],
                        row["role"], row["is_active"], row.get("is_capture_manager", 0))
            login_user(user, remember=False)  # Sessions expire when browser closes
            db.reset_failed_login(row["id"])
            db.record_login(row["id"], ip)
            db.write_audit_log("LOGIN_SUCCESS", username=username, user_id=row["id"],
                               ip_address=ip)
            # Admins land on the user management panel
            if user.is_admin:
                return redirect(url_for("admin_users"))
            return redirect(request.args.get("next") or url_for("dashboard"))
        else:
            threshold = int(db.get_app_setting("lockout_threshold", "5"))
            just_locked = db.increment_failed_login(row["id"], threshold)
            if just_locked:
                db.write_audit_log("ACCOUNT_LOCKED", username=username, user_id=row["id"],
                                   detail=f"Locked after {threshold} failed attempts",
                                   ip_address=ip)
                flash("Too many failed attempts. Account has been locked — contact an admin.", "danger")
            else:
                db.write_audit_log("LOGIN_FAILED", username=username, user_id=row["id"],
                                   detail="Invalid password", ip_address=ip)
                flash("Invalid username or password.", "danger")
    return render_template("login.html")


@app.route("/forgot-username", methods=["POST"])
def forgot_username():
    """Look up username by email address and display it on-screen."""
    email = request.form.get("email", "").strip()
    if not email:
        flash("Please enter your email address.", "warning")
        return redirect(url_for("login"))

    user = db.get_user_by_email(email)
    ip   = request.remote_addr or "unknown"
    if user:
        db.write_audit_log("FORGOT_USERNAME", username=user["username"],
                           user_id=user["id"], detail="Username lookup by email",
                           ip_address=ip)
        flash(f"Your username is: <strong>{user['username']}</strong>", "info")
    else:
        # Don't reveal whether email exists — same message either way
        flash("If that email is registered, your username has been displayed above.", "info")

    return redirect(url_for("login"))


@app.route("/forgot-password", methods=["POST"])
def forgot_password():
    """Generate a reset token and email it to the user."""
    email = request.form.get("email", "").strip()
    ip    = request.remote_addr or "unknown"

    if not email:
        flash("Please enter your email address.", "warning")
        return redirect(url_for("login"))

    user = db.get_user_by_email(email)
    if user and user["is_active"]:
        token = db.create_password_reset_token(user["id"], hours_valid=2)
        reset_url = url_for("reset_password", token=token, _external=True)

        subject   = "CaptureIQ — Password Reset Request"
        body_text = (
            f"Hi {user['username']},\n\n"
            f"A password reset was requested for your CaptureIQ account.\n\n"
            f"Click the link below to reset your password (valid for 2 hours):\n"
            f"{reset_url}\n\n"
            f"If you did not request this, you can safely ignore this email.\n\n"
            f"— CaptureIQ"
        )
        body_html = f"""
        <div style="font-family:system-ui,sans-serif;max-width:480px;margin:0 auto">
          <div style="background:#003087;color:#fff;padding:1.5rem;border-radius:8px 8px 0 0;text-align:center">
            <h2 style="margin:0">🔑 Password Reset</h2>
            <p style="margin:.5rem 0 0;opacity:.8">CaptureIQ</p>
          </div>
          <div style="background:#fff;padding:2rem;border:1px solid #dee2e6;border-top:none;border-radius:0 0 8px 8px">
            <p>Hi <strong>{user['username']}</strong>,</p>
            <p>A password reset was requested for your account. Click the button below to choose a new password.</p>
            <p style="text-align:center;margin:2rem 0">
              <a href="{reset_url}"
                 style="background:#003087;color:#fff;padding:.75rem 2rem;border-radius:6px;
                        text-decoration:none;font-weight:600;display:inline-block">
                Reset My Password
              </a>
            </p>
            <p style="color:#6c757d;font-size:.85rem">
              This link expires in <strong>2 hours</strong>. If you did not request a reset, ignore this email.
            </p>
          </div>
        </div>
        """

        ok, err = db.send_email(email, subject, body_text, body_html)
        db.write_audit_log("PASSWORD_RESET_REQUESTED", username=user["username"],
                           user_id=user["id"],
                           detail=f"Reset email {'sent' if ok else 'failed: ' + err}",
                           ip_address=ip)

        if not ok:
            # Email not configured — show token info to admin in flash for local installs
            smtp_host = db.get_app_setting("smtp_host", "")
            if not smtp_host:
                flash(
                    f"Email is not configured. An admin can reset the password manually using "
                    f"<code>python3 reset_password.py</code>, or configure SMTP in App Settings.",
                    "warning"
                )
            else:
                flash("Failed to send reset email. Contact your administrator.", "danger")
            return redirect(url_for("login"))

    # Always show the same message whether email matched or not
    flash("If that email address is registered, a password reset link has been sent. Check your inbox.", "info")
    return redirect(url_for("login"))


@app.route("/reset-password/<token>", methods=["GET", "POST"])
def reset_password(token: str):
    """Validate token and allow user to set a new password."""
    row = db.get_valid_reset_token(token)
    if not row:
        flash("This password reset link is invalid or has expired. Please request a new one.", "danger")
        return redirect(url_for("login"))

    if request.method == "POST":
        password  = request.form.get("password", "")
        confirm   = request.form.get("confirm", "")
        min_len   = int(db.get_app_setting("min_password_length", "8"))

        if len(password) < min_len:
            flash(f"Password must be at least {min_len} characters.", "warning")
            return render_template("reset_password.html", token=token, username=row["username"])
        if password != confirm:
            flash("Passwords do not match.", "warning")
            return render_template("reset_password.html", token=token, username=row["username"])

        from werkzeug.security import generate_password_hash
        new_hash = generate_password_hash(password)
        ok = db.consume_reset_token(token, new_hash)
        if ok:
            db.write_audit_log("PASSWORD_RESET_COMPLETED", username=row["username"],
                               user_id=row["user_id"],
                               ip_address=request.remote_addr or "unknown")
            flash("Your password has been reset. Please sign in.", "success")
        else:
            flash("Reset link expired. Please request a new one.", "danger")
        return redirect(url_for("login"))

    return render_template("reset_password.html", token=token, username=row["username"])


@app.route("/register", methods=["GET", "POST"])
def register():
    """Self-service account registration. Requires admin approval if that setting is on."""
    if current_user.is_authenticated:
        return redirect(url_for("dashboard"))
    if db.count_users() == 0:
        return redirect(url_for("setup"))

    # Check if registration is enabled
    allow_reg = db.get_app_setting("allow_registration", "true").lower() == "true"
    if not allow_reg:
        flash("Self-registration is currently disabled. Contact an administrator.", "warning")
        return redirect(url_for("login"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email    = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        confirm  = request.form.get("confirm", "")
        ip       = request.remote_addr or "unknown"
        min_len  = int(db.get_app_setting("min_password_length", "8"))

        error = None
        if not username or not password:
            error = "Username and password are required."
        elif len(username) < 3:
            error = "Username must be at least 3 characters."
        elif db.get_user_by_username(username):
            error = f"Username '{username}' is already taken."
        elif email and db.get_user_by_email(email):
            error = "An account with that email already exists."
        elif len(password) < min_len:
            error = f"Password must be at least {min_len} characters."
        elif password != confirm:
            error = "Passwords do not match."

        if error:
            flash(error, "danger")
            return render_template("register.html", username=username, email=email)

        needs_approval = db.get_app_setting("registration_approval", "false").lower() == "true"
        if needs_approval:
            # Create account as inactive — admin must approve
            db.create_user(username, email, generate_password_hash(password),
                           role="user", is_active=0)
            db.write_audit_log("REGISTRATION_PENDING", username=username,
                               detail="Self-registration pending admin approval",
                               ip_address=ip)
            flash("Your account request has been submitted. An administrator will review and activate it shortly.", "info")
        else:
            db.create_user(username, email, generate_password_hash(password), role="user")
            db.write_audit_log("REGISTRATION_COMPLETED", username=username,
                               detail="Self-registration completed (auto-approved)",
                               ip_address=ip)
            flash("Account created! You can now sign in.", "success")

        return redirect(url_for("login"))

    return render_template("register.html", username="", email="")


@app.route("/admin/registrations")
@login_required
def admin_pending_registrations():
    """List inactive users who registered and are awaiting approval."""
    if not current_user.is_admin:
        return redirect(url_for("admin_users"))
    pending = [u for u in db.get_all_users() if not u["is_active"] and u["role"] == "user"]
    return render_template("admin_users.html",
                           users=db.get_all_users(),
                           pending_registrations=pending)


@app.route("/admin/users/<int:user_id>/approve", methods=["POST"])
@login_required
def admin_approve_registration(user_id: int):
    if not current_user.is_admin:
        abort(403)
    with db.get_db() as conn:
        conn.execute("UPDATE users SET is_active=1 WHERE id=?", (user_id,))
    u = db.get_user_by_id(user_id)
    db.write_audit_log("USER_APPROVED", username=current_user.username,
                       user_id=current_user.id,
                       detail=f"Approved registration for '{u['username'] if u else user_id}'",
                       ip_address=request.remote_addr)
    flash("Account approved and activated.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/reject", methods=["POST"])
@login_required
def admin_reject_registration(user_id: int):
    if not current_user.is_admin:
        abort(403)
    u = db.get_user_by_id(user_id)
    db.write_audit_log("USER_REJECTED", username=current_user.username,
                       user_id=current_user.id,
                       detail=f"Rejected registration for '{u['username'] if u else user_id}'",
                       ip_address=request.remote_addr)
    with db.get_db() as conn:
        conn.execute("DELETE FROM users WHERE id=? AND is_active=0", (user_id,))
    flash("Registration request rejected and removed.", "info")
    return redirect(url_for("admin_users"))


@app.route("/logout")
@login_required
def logout():
    db.write_audit_log("LOGOUT", username=current_user.username,
                       user_id=current_user.id,
                       ip_address=request.remote_addr or "unknown")
    logout_user()
    flash("You have been logged out.", "info")
    return redirect(url_for("login"))


@app.route("/setup", methods=["GET", "POST"])
def setup():
    """First-run admin account creation. Only accessible when no users exist."""
    if db.count_users() > 0:
        return redirect(url_for("login"))
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        email    = request.form.get("email", "").strip()
        password = request.form.get("password", "")
        confirm  = request.form.get("confirm", "")
        if not username or not password:
            flash("Username and password are required.", "danger")
        elif password != confirm:
            flash("Passwords do not match.", "danger")
        elif len(password) < 6:
            flash("Password must be at least 6 characters.", "danger")
        else:
            db.create_user(username, email, generate_password_hash(password), role="admin")
            flash(f"Admin account '{username}' created. Please log in.", "success")
            return redirect(url_for("login"))
    return render_template("setup.html")


# ── Admin — User Management ────────────────────────────────────────────────────

@app.route("/admin/users")
@login_required
def admin_users():
    if not current_user.is_admin:
        flash("Admin access required.", "danger")
        return redirect(url_for("dashboard"))
    users = db.get_all_users()
    return render_template("admin_users.html", users=users)


@app.route("/admin/users/create", methods=["POST"])
@login_required
def admin_create_user():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    username = request.form.get("username", "").strip()
    email    = request.form.get("email", "").strip()
    password = request.form.get("password", "")
    role     = request.form.get("role", "user")
    if not username or not password:
        flash("Username and password are required.", "danger")
    elif db.get_user_by_username(username):
        flash(f"Username '{username}' is already taken.", "danger")
    elif len(password) < 6:
        flash("Password must be at least 6 characters.", "danger")
    else:
        actual_role = "user" if role == "capture_manager" else role
        new_user_id = db.create_user(username, email, generate_password_hash(password), actual_role)
        # If capture_manager role selected, elevate the user
        if role == "capture_manager" and new_user_id:
            with db.get_db() as conn:
                conn.execute("UPDATE users SET is_capture_manager=1 WHERE id=?", (new_user_id,))
        db.write_audit_log("USER_CREATED",
                           username=current_user.username, user_id=current_user.id,
                           detail=f"Created user '{username}' with role '{role}'",
                           ip_address=request.remote_addr)
        flash(f"User '{username}' created successfully.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/toggle", methods=["POST"])
@login_required
def admin_toggle_user(user_id):
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    if user_id == current_user.id:
        flash("You cannot deactivate your own account.", "warning")
    else:
        row = db.get_user_by_id(user_id)
        if row:
            new_state = not row["is_active"]
            db.set_user_active(user_id, new_state)
            action = "USER_REACTIVATED" if new_state else "USER_DEACTIVATED"
            db.write_audit_log(action,
                               username=current_user.username, user_id=current_user.id,
                               detail=f"{'Activated' if new_state else 'Deactivated'} user '{row['username']}'",
                               ip_address=request.remote_addr)
            flash(f"User '{row['username']}' {'activated' if new_state else 'deactivated'}.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/unlock", methods=["POST"])
@login_required
def admin_unlock_user(user_id):
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    row = db.get_user_by_id(user_id)
    if row:
        db.unlock_user(user_id)
        db.write_audit_log("ACCOUNT_UNLOCKED",
                           username=current_user.username, user_id=current_user.id,
                           detail=f"Unlocked account for '{row['username']}'",
                           ip_address=request.remote_addr)
        flash(f"Account for '{row['username']}' has been unlocked.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/delete", methods=["POST"])
@login_required
def admin_delete_user(user_id):
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    if user_id == current_user.id:
        flash("You cannot delete your own account.", "warning")
    else:
        row = db.get_user_by_id(user_id)
        if row:
            db.write_audit_log("USER_DELETED",
                               username=current_user.username, user_id=current_user.id,
                               detail=f"Deleted user '{row['username']}'",
                               ip_address=request.remote_addr)
            db.delete_user(user_id)
            flash(f"User '{row['username']}' deleted.", "success")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/reset-password", methods=["POST"])
@login_required
def admin_reset_password(user_id):
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    new_password = request.form.get("new_password", "")
    if len(new_password) < 6:
        flash("Password must be at least 6 characters.", "danger")
    else:
        db.update_user_password(user_id, generate_password_hash(new_password))
        row = db.get_user_by_id(user_id)
        db.write_audit_log("USER_PASSWORD_RESET",
                           username=current_user.username, user_id=current_user.id,
                           detail=f"Reset password for '{row['username']}'",
                           ip_address=request.remote_addr)
        flash(f"Password reset for '{row['username']}'.", "success")
    return redirect(url_for("admin_users"))


# ── Admin — Capture Manager Role Management ────────────────────────────────────

@app.route("/admin/users/<int:user_id>/elevate-capture-manager", methods=["POST"])
@login_required
def elevate_capture_manager(user_id):
    """Elevate a user to capture manager role (admin only)."""
    if not current_user.is_admin:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Admin access required'}), 403
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for("dashboard"))

    user_row = db.get_user_by_id(user_id)
    if not user_row:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'User not found'}), 404
        abort(404)

    try:
        # Get a database connection
        conn = db.get_db()

        # First try with is_capture_manager, if it fails, add the column
        try:
            conn.execute("UPDATE users SET is_capture_manager = 1 WHERE id = ?", (user_id,))
            conn.commit()
        except Exception as e:
            if "no such column" in str(e):
                # Column doesn't exist, add it
                try:
                    conn.execute("ALTER TABLE users ADD COLUMN is_capture_manager INTEGER DEFAULT 0")
                    conn.commit()
                except:
                    pass  # Column may already exist from concurrent operation

                conn.execute("UPDATE users SET is_capture_manager = 1 WHERE id = ?", (user_id,))
                conn.commit()
            else:
                raise

        # Log the change
        reason = None
        if request.is_json:
            reason = request.json.get('reason', '') if request.json else ''
        else:
            reason = request.form.get('reason', '')

        conn.execute(
            """INSERT INTO role_change_history
               (user_id, role_changed_to, changed_by_user_id, reason)
               VALUES (?, ?, ?, ?)""",
            (user_id, 'capture_manager', current_user.id, reason or 'Elevated to Capture Manager')
        )
        conn.commit()
        conn.close()

        db.write_audit_log("CAPTURE_MANAGER_ELEVATED",
                           username=current_user.username, user_id=current_user.id,
                           detail=f"Elevated '{user_row['username']}' to Capture Manager",
                           ip_address=request.remote_addr)

        if request.is_json:
            return jsonify({'status': 'success', 'message': f'{user_row["username"]} is now a Capture Manager'})
        else:
            flash(f"{user_row['username']} is now a Capture Manager.", "success")
            return redirect(url_for('admin_users'))

    except Exception as e:
        print(f"DEBUG: Elevation error: {str(e)}")
        import traceback
        traceback.print_exc()
        if request.is_json:
            return jsonify({'status': 'error', 'message': str(e)}), 500
        else:
            flash(f"Error elevating user: {str(e)}", "danger")
            return redirect(url_for('admin_users'))


@app.route("/admin/users/<int:user_id>/remove-capture-manager", methods=["POST"])
@login_required
def remove_capture_manager(user_id):
    """Remove capture manager role from a user (admin only)."""
    if not current_user.is_admin:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Admin access required'}), 403
        flash("You do not have permission to perform this action.", "danger")
        return redirect(url_for("dashboard"))

    user_row = db.get_user_by_id(user_id)
    if not user_row:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'User not found'}), 404
        abort(404)

    try:
        # Get a database connection
        conn = db.get_db()

        # First try with is_capture_manager, if it fails, add the column
        try:
            conn.execute("UPDATE users SET is_capture_manager = 0 WHERE id = ?", (user_id,))
            conn.commit()
        except Exception as e:
            if "no such column" in str(e):
                # Column doesn't exist, add it
                try:
                    conn.execute("ALTER TABLE users ADD COLUMN is_capture_manager INTEGER DEFAULT 0")
                    conn.commit()
                except:
                    pass  # Column may already exist from concurrent operation

                conn.execute("UPDATE users SET is_capture_manager = 0 WHERE id = ?", (user_id,))
                conn.commit()
            else:
                raise

        # Log the change
        conn.execute(
            """INSERT INTO role_change_history
               (user_id, role_changed_to, changed_by_user_id, reason)
               VALUES (?, ?, ?, ?)""",
            (user_id, 'user', current_user.id, 'Capture manager role removed')
        )
        conn.commit()
        conn.close()

        db.write_audit_log("CAPTURE_MANAGER_REMOVED",
                           username=current_user.username, user_id=current_user.id,
                           detail=f"Removed Capture Manager role from '{user_row['username']}'",
                           ip_address=request.remote_addr)

        if request.is_json:
            return jsonify({'status': 'success', 'message': f'{user_row["username"]} is no longer a Capture Manager'})
        else:
            flash(f"{user_row['username']} is no longer a Capture Manager.", "success")
            return redirect(url_for('admin_users'))

    except Exception as e:
        if request.is_json:
            return jsonify({'status': 'error', 'message': str(e)}), 500
        else:
            flash(f"Error removing capture manager: {str(e)}", "danger")
            return redirect(url_for('admin_users'))


# ── Admin — Audit Log ──────────────────────────────────────────────────────────

@app.route("/admin/audit-log")
@login_required
def admin_audit_log():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    page        = int(request.args.get("page", 1))
    action_type = request.args.get("action_type", "")
    username    = request.args.get("username", "")
    per_page    = 50
    result = db.get_audit_log(
        page=page, per_page=per_page,
        action_type=action_type or None,
        username=username or None,
    )
    action_types = [
        "LOGIN_SUCCESS", "LOGIN_FAILED", "LOGOUT",
        "ACCOUNT_LOCKED", "ACCOUNT_UNLOCKED",
        "USER_CREATED", "USER_DELETED", "USER_DEACTIVATED",
        "USER_REACTIVATED", "USER_PASSWORD_RESET",
        "INGEST_STARTED", "INGEST_COMPLETED", "INGEST_FAILED",
    ]
    return render_template("admin_audit_log.html",
                           result=result, action_types=action_types,
                           action_type=action_type, username=username)


@app.route("/admin/audit-log/export")
@login_required
def admin_audit_log_export():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    rows = db.get_audit_log_csv()
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["ID", "Timestamp", "Username", "User ID", "Action", "Detail", "IP Address"])
    for r in rows:
        writer.writerow([r["id"], r["timestamp"], r["username"], r["user_id"],
                         r["action_type"], r["detail"], r["ip_address"]])
    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=captureiq_audit_log.csv"},
    )


# ── Admin — Database & Backup ──────────────────────────────────────────────────

@app.route("/admin/db")
@login_required
def admin_db():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    stats = db.get_db_stats()
    backup_log = db.get_backup_log()
    return render_template("admin_db.html", stats=stats, backup_log=backup_log)


@app.route("/admin/db/backup")
@login_required
def admin_db_backup():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    filename = f"captureiq_backup_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.db"
    db.write_audit_log("DB_BACKUP",
                       username=current_user.username, user_id=current_user.id,
                       detail="Manual database backup downloaded",
                       ip_address=request.remote_addr)
    import os as _os
    db.log_db_backup(filename=filename, size_bytes=_os.path.getsize(db.DB_PATH), user_id=current_user.id)
    return send_file(db.DB_PATH, as_attachment=True, download_name=filename)


@app.route("/admin/db/import", methods=["POST"])
@login_required
def admin_db_import():
    if not current_user.is_admin:
        abort(403)
    f = request.files.get("db_file")
    if not f or not f.filename.endswith(".db"):
        flash("Please upload a valid .db file.", "danger")
        return redirect(url_for("admin_db"))
    import shutil as _shutil, os as _os
    db_path = db.DB_PATH
    backup_path = db_path + ".pre_import_backup"
    _shutil.copy2(db_path, backup_path)
    try:
        f.save(db_path)
        db.write_audit_log("DB_IMPORTED", username=current_user.username,
                          user_id=current_user.id,
                          detail="Database restored from uploaded file",
                          ip_address=request.remote_addr)
        flash("Database imported successfully. A backup of the previous database was saved.", "success")
    except Exception as e:
        _shutil.copy2(backup_path, db_path)
        flash(f"Import failed: {str(e)}. Previous database restored.", "danger")
    return redirect(url_for("admin_db"))


# ── Admin — Settings ───────────────────────────────────────────────────────────

@app.route("/admin/settings")
@login_required
def admin_settings():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    settings = {
        "lockout_threshold":    db.get_app_setting("lockout_threshold", "5"),
        "min_password_length":  db.get_app_setting("min_password_length", "8"),
        "allow_registration":   db.get_app_setting("allow_registration", "true"),
        "registration_approval": db.get_app_setting("registration_approval", "false"),
        "smtp_host":     db.get_app_setting("smtp_host", ""),
        "smtp_port":     db.get_app_setting("smtp_port", "587"),
        "smtp_user":     db.get_app_setting("smtp_user", ""),
        "smtp_password": db.get_app_setting("smtp_password", ""),
        "smtp_from":     db.get_app_setting("smtp_from", ""),
        "smtp_tls":      db.get_app_setting("smtp_tls", "true"),
        "smtp_test_result": None,
        "smtp_test_ok": False,
    }
    return render_template("admin_settings.html", settings=settings)


@app.route("/admin/settings/save", methods=["POST"])
@login_required
def admin_save_settings():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))
    lockout_threshold  = request.form.get("lockout_threshold", "5").strip()
    min_password_length = request.form.get("min_password_length", "6").strip()
    try:
        lt = int(lockout_threshold)
        if lt < 1:
            raise ValueError
    except ValueError:
        flash("Lockout threshold must be a positive integer.", "danger")
        return redirect(url_for("admin_settings"))
    try:
        mpl = int(min_password_length)
        if mpl < 4:
            raise ValueError
    except ValueError:
        flash("Minimum password length must be at least 4.", "danger")
        return redirect(url_for("admin_settings"))
    allow_reg  = "true" if request.form.get("allow_registration") else "false"
    req_approval = "true" if request.form.get("registration_approval") else "false"
    db.set_app_setting("lockout_threshold", str(lt))
    db.set_app_setting("min_password_length", str(mpl))
    db.set_app_setting("allow_registration", allow_reg)
    db.set_app_setting("registration_approval", req_approval)
    db.write_audit_log("SETTINGS_CHANGED",
                       username=current_user.username, user_id=current_user.id,
                       detail=f"lockout_threshold={lt}, min_password_length={mpl}, "
                              f"allow_registration={allow_reg}, registration_approval={req_approval}",
                       ip_address=request.remote_addr)
    flash("Settings saved.", "success")
    return redirect(url_for("admin_settings"))


@app.route("/admin/settings/smtp", methods=["POST"])
@login_required
def admin_save_smtp():
    if not current_user.is_admin:
        return redirect(url_for("dashboard"))

    smtp_host = request.form.get("smtp_host", "").strip()
    smtp_port = request.form.get("smtp_port", "587").strip() or "587"
    smtp_user = request.form.get("smtp_user", "").strip()
    smtp_pass = request.form.get("smtp_password", "").strip()
    smtp_from = request.form.get("smtp_from", "").strip()
    smtp_tls  = "true" if request.form.get("smtp_tls") else "false"

    db.set_app_setting("smtp_host", smtp_host)
    db.set_app_setting("smtp_port", smtp_port)
    db.set_app_setting("smtp_user", smtp_user)
    if smtp_pass:  # Only update password if a new one was entered
        db.set_app_setting("smtp_password", smtp_pass)
    db.set_app_setting("smtp_from", smtp_from)
    db.set_app_setting("smtp_tls", smtp_tls)

    send_test = request.form.get("send_test")
    test_result = None
    test_ok     = False

    if send_test and smtp_host and smtp_user:
        test_to = smtp_from or smtp_user
        ok, err = db.send_email(
            test_to,
            "CaptureIQ — SMTP Test",
            "This is a test email from CaptureIQ to confirm your SMTP settings are working.",
        )
        test_ok     = ok
        test_result = f"Test email sent to {test_to}." if ok else f"Failed: {err}"
        db.write_audit_log("SMTP_TEST", username=current_user.username,
                           user_id=current_user.id,
                           detail=f"SMTP test {'OK' if ok else 'FAILED: ' + err}",
                           ip_address=request.remote_addr)
    else:
        flash("Email settings saved.", "success")

    settings = {
        "lockout_threshold":  db.get_app_setting("lockout_threshold", "5"),
        "min_password_length": db.get_app_setting("min_password_length", "6"),
        "smtp_host":     smtp_host,
        "smtp_port":     smtp_port,
        "smtp_user":     smtp_user,
        "smtp_password": db.get_app_setting("smtp_password", ""),
        "smtp_from":     smtp_from,
        "smtp_tls":      smtp_tls,
        "smtp_test_result": test_result,
        "smtp_test_ok":     test_ok,
    }
    return render_template("admin_settings.html", settings=settings)


@app.route("/help/user-guide")
@app.route("/guide")
@login_required
def user_guide():
    """Display the user guide."""
    return render_template("user_guide.html")


# ── Dashboard ──────────────────────────────────────────────────────────────────

@app.route("/")
@login_required
def dashboard():
    stats = db.get_stats(user_id=current_user.id)
    capture_stats = db.get_capture_stats()
    pending_invitations = db.get_pending_invitations(current_user.id)
    # Compute total pipeline value from active capture plans
    pipeline_value = 0
    try:
        with db.get_db() as conn:
            row = conn.execute("""
                SELECT COALESCE(SUM(target_contract_value),0) as total
                FROM capture_plans WHERE is_archived=0 OR is_archived IS NULL
            """).fetchone()
            pipeline_value = row['total'] if row else 0
    except Exception:
        pipeline_value = 0
    pipeline_value_fmt = "{:,.0f}".format(pipeline_value)
    return render_template("dashboard.html", stats=stats, jobs=_jobs,
                           capture_stats=capture_stats,
                           pending_invitations=pending_invitations,
                           pipeline_value=pipeline_value,
                           pipeline_value_fmt=pipeline_value_fmt)


# ── Capture Management ─────────────────────────────────────────────────────────

@app.route("/capture")
@login_required
@require_capture_manager
def capture_dashboard():
    """Capture manager dashboard - list all capture plans grouped by stage."""
    plans = db.get_capture_plans_by_user(current_user.id, include_archived=False)

    # Group by stage
    by_stage = {}
    for plan in plans:
        stage = plan['stage'] or 'pre-release'
        if stage not in by_stage:
            by_stage[stage] = []
        by_stage[stage].append(plan)

    return render_template("capture/dashboard.html", plans_by_stage=by_stage)


@app.route("/capture/create", methods=["GET", "POST"])
@login_required
@require_capture_manager
def create_capture_plan():
    """Create a new capture plan."""
    if request.method == "GET":
        # Get available topics for linking
        topics = db.get_topics(limit=1000)
        return render_template("capture/create_plan.html", topics=topics)

    if request.method == "POST":
        data = request.form if request.form else request.json

        try:
            plan_id = db.create_capture_plan(
                capture_name=data.get("capture_name"),
                capture_lead_id=current_user.id,
                created_by_user_id=current_user.id,
                solicitation_id=data.get("solicitation_id") or None,
                customer_name=data.get("customer_name"),
                customer_website=data.get("customer_website"),
                estimated_release_date=data.get("estimated_release_date"),
                proposal_due_date=data.get("proposal_due_date"),
                target_contract_value=float(data.get("target_contract_value")) if data.get("target_contract_value") else None,
                stage=data.get("stage", "pre-release"),
                confidence_level=data.get("confidence_level", "medium"),
                win_probability=int(data.get("win_probability", 50))
            )

            # Add creator as owner in access table
            db.add_capture_plan_access(plan_id, current_user.id, "owner")

            db.write_audit_log("CAPTURE_PLAN_CREATED",
                             username=current_user.username, user_id=current_user.id,
                             detail=f"Created capture plan '{data.get('capture_name')}'",
                             ip_address=request.remote_addr)

            flash(f"Capture plan created successfully.", "success")

            if request.is_json:
                return jsonify({"status": "success", "capture_plan_id": plan_id})
            else:
                return redirect(url_for("capture_plan_detail", plan_id=plan_id))

        except Exception as e:
            error_msg = f"Failed to create capture plan: {str(e)}"
            if request.is_json:
                return jsonify({"status": "error", "message": error_msg}), 400
            else:
                flash(error_msg, "danger")
                return redirect(url_for("create_capture_plan"))


@app.route("/capture/<int:plan_id>")
@login_required
@require_capture_manager
def capture_plan_detail(plan_id):
    """View a capture plan."""
    plan = db.get_capture_plan(plan_id)

    if not plan:
        abort(404)

    # Check access: must be lead, have explicit access, or be admin
    if plan['capture_lead_id'] != current_user.id and not current_user.is_admin:
        access = db.get_capture_plan_access(plan_id, current_user.id)
        if not access:
            abort(403)

    # Get linked projects
    linked_projects = db.get_projects_by_capture_plan(plan_id)

    # Get team members
    members = db.list_capture_plan_members(plan_id)

    # Get solicitation details if linked
    solicitation = None
    if plan['solicitation_id']:
        solicitation = db.get_topic(plan['solicitation_id'])

    # Populate Add Member modal
    all_users = db.get_all_users()
    member_ids = {m['id'] for m in members}
    available_members = [u for u in all_users if u['is_active'] and u['id'] not in member_ids]

    # Populate Link Project modal
    all_projects = db.get_projects_by_owner(current_user.id)
    linked_ids = {p['id'] for p in linked_projects}
    available_projects = [p for p in all_projects if p['id'] not in linked_ids]

    # Get key contacts
    key_contacts = db.get_key_contacts(plan_id)

    return render_template("capture/plan_detail.html",
                         plan=plan,
                         solicitation=solicitation,
                         linked_projects=linked_projects,
                         members=members,
                         is_owner=plan['capture_lead_id'] == current_user.id or current_user.is_admin,
                         available_members=available_members,
                         available_projects=available_projects,
                         key_contacts=key_contacts)


@app.route("/capture/<int:plan_id>/edit", methods=["POST"])
@login_required
@require_capture_manager
def edit_capture_plan(plan_id):
    """Update a capture plan."""
    plan = db.get_capture_plan(plan_id)

    if not plan:
        abort(404)

    # Only owner or admin can edit
    if plan['capture_lead_id'] != current_user.id and not current_user.is_admin:
        abort(403)

    data = request.form if request.form else request.json

    try:
        db.update_capture_plan(
            plan_id,
            capture_name=data.get("capture_name"),
            customer_name=data.get("customer_name"),
            customer_website=data.get("customer_website"),
            estimated_release_date=data.get("estimated_release_date"),
            proposal_due_date=data.get("proposal_due_date"),
            target_contract_value=float(data.get("target_contract_value")) if data.get("target_contract_value") else None,
            stage=data.get("stage"),
            confidence_level=data.get("confidence_level"),
            win_probability=int(data.get("win_probability")) if data.get("win_probability") else None
        )

        db.write_audit_log("CAPTURE_PLAN_UPDATED",
                         username=current_user.username, user_id=current_user.id,
                         detail=f"Updated capture plan '{plan['capture_name']}'",
                         ip_address=request.remote_addr)

        flash("Capture plan updated successfully.", "success")

        if request.is_json:
            return jsonify({"status": "success"})
        else:
            return redirect(url_for("capture_plan_detail", plan_id=plan_id))

    except Exception as e:
        error_msg = f"Failed to update capture plan: {str(e)}"
        if request.is_json:
            return jsonify({"status": "error", "message": error_msg}), 400
        else:
            flash(error_msg, "danger")
            return redirect(url_for("capture_plan_detail", plan_id=plan_id))


@app.route("/capture/<int:plan_id>/archive", methods=["POST"])
@login_required
@require_capture_manager
def archive_capture_plan(plan_id):
    """Archive a capture plan."""
    plan = db.get_capture_plan(plan_id)

    if not plan:
        abort(404)

    # Only owner or admin can archive
    if plan['capture_lead_id'] != current_user.id and not current_user.is_admin:
        abort(403)

    db.update_capture_plan(plan_id, is_archived=1)

    db.write_audit_log("CAPTURE_PLAN_ARCHIVED",
                     username=current_user.username, user_id=current_user.id,
                     detail=f"Archived capture plan '{plan['capture_name']}'",
                     ip_address=request.remote_addr)

    flash("Capture plan archived.", "success")
    return redirect(url_for("capture_dashboard"))


@app.route("/capture/<int:plan_id>/add-member", methods=["POST"])
@login_required
@require_capture_manager
def add_capture_plan_member(plan_id):
    """Add a team member to a capture plan."""
    plan = db.get_capture_plan(plan_id)

    if not plan:
        abort(404)

    # Only owner or admin can add members
    if plan['capture_lead_id'] != current_user.id and not current_user.is_admin:
        abort(403)

    data = request.form if request.form else request.json
    user_id = data.get("user_id")
    access_level = data.get("access_level", "viewer")

    user = db.get_user_by_id(user_id)
    if not user:
        return jsonify({"status": "error", "message": "User not found"}), 404

    if db.add_capture_plan_access(plan_id, user_id, access_level):
        db.write_audit_log("CAPTURE_PLAN_MEMBER_ADDED",
                         username=current_user.username, user_id=current_user.id,
                         detail=f"Added {user['username']} to capture plan '{plan['capture_name']}'",
                         ip_address=request.remote_addr)

        flash(f"{user['username']} added to capture plan.", "success")

    if request.is_json:
        return jsonify({"status": "success"})
    else:
        return redirect(url_for("capture_plan_detail", plan_id=plan_id))


@app.route("/capture/<int:plan_id>/remove-member/<int:member_id>", methods=["POST"])
@login_required
@require_capture_manager
def remove_capture_plan_member(plan_id, member_id):
    """Remove a team member from a capture plan."""
    plan = db.get_capture_plan(plan_id)

    if not plan:
        abort(404)

    # Only owner or admin can remove members
    if plan['capture_lead_id'] != current_user.id and not current_user.is_admin:
        abort(403)

    # Can't remove the owner
    if member_id == plan['capture_lead_id']:
        return jsonify({"status": "error", "message": "Cannot remove the capture lead"}), 403

    user = db.get_user_by_id(member_id)

    db.remove_capture_plan_access(plan_id, member_id)

    db.write_audit_log("CAPTURE_PLAN_MEMBER_REMOVED",
                     username=current_user.username, user_id=current_user.id,
                     detail=f"Removed {user['username']} from capture plan '{plan['capture_name']}'",
                     ip_address=request.remote_addr)

    flash(f"{user['username']} removed from capture plan.", "success")

    if request.is_json:
        return jsonify({"status": "success"})
    else:
        return redirect(url_for("capture_plan_detail", plan_id=plan_id))


@app.route("/capture/<int:plan_id>/link-project", methods=["POST"])
@login_required
@require_capture_manager
def link_project_from_capture(plan_id):
    plan = db.get_capture_plan(plan_id)
    if not plan:
        abort(404)
    if plan['capture_lead_id'] != current_user.id and not current_user.is_admin:
        abort(403)
    project_id = request.form.get("project_id", type=int)
    if project_id:
        db.link_project_to_capture_plan(project_id, plan_id)
        flash("Project linked successfully.", "success")
    return redirect(url_for("capture_plan_detail", plan_id=plan_id))


@app.route("/capture/<int:plan_id>/contacts/add", methods=["POST"])
@login_required
@require_capture_manager
def add_key_contact(plan_id):
    plan = db.get_capture_plan(plan_id)
    if not plan:
        abort(404)
    db.add_key_contact(
        capture_plan_id=plan_id,
        first_name=request.form.get("first_name", "").strip(),
        last_name=request.form.get("last_name", "").strip(),
        email=request.form.get("email", "").strip() or None,
        phone=request.form.get("phone", "").strip() or None,
        agency=request.form.get("agency", "").strip() or None,
        title=request.form.get("title", "").strip() or None,
        notes=request.form.get("notes", "").strip() or None,
        last_contacted=request.form.get("last_contacted") or None,
        record_owner_id=current_user.id
    )
    flash("Contact added.", "success")
    return redirect(url_for("capture_plan_detail", plan_id=plan_id) + "#contacts")


@app.route("/capture/contacts/<int:contact_id>/delete", methods=["POST"])
@login_required
@require_capture_manager
def delete_key_contact(contact_id):
    with db.get_db() as conn:
        row = conn.execute("SELECT capture_plan_id FROM key_contacts WHERE id=?", (contact_id,)).fetchone()
    plan_id = row['capture_plan_id'] if row else None
    db.delete_key_contact(contact_id)
    flash("Contact removed.", "info")
    if plan_id:
        return redirect(url_for("capture_plan_detail", plan_id=plan_id) + "#contacts")
    return redirect(url_for("capture_dashboard"))


@app.route("/contacts")
@login_required
@require_capture_manager
def all_key_contacts():
    search = request.args.get("q", "").strip()
    agency = request.args.get("agency", "").strip()
    sort = request.args.get("sort", "last_name")
    contacts = db.get_key_contacts()
    if search:
        sl = search.lower()
        contacts = [c for c in contacts if sl in (c.get('first_name') or '').lower()
                    or sl in (c.get('last_name') or '').lower()
                    or sl in (c.get('agency') or '').lower()
                    or sl in (c.get('email') or '').lower()]
    if agency:
        contacts = [c for c in contacts if agency.lower() in (c.get('agency') or '').lower()]
    valid_sorts = ['last_name', 'first_name', 'agency', 'record_owner_name']
    if sort not in valid_sorts:
        sort = 'last_name'
    contacts.sort(key=lambda c: (c.get(sort) or '').lower())
    agencies = sorted({c['agency'] for c in db.get_key_contacts() if c.get('agency')})
    return render_template("contacts.html", contacts=contacts, agencies=agencies,
                           search=search, agency=agency, sort=sort)


@app.route("/project/<int:project_id>/link-capture-plan", methods=["POST"])
@login_required
def link_project_to_capture(project_id):
    """Link a project to a capture plan."""
    project = db.get_project(project_id)

    if not project:
        abort(404)

    # User must have edit access to project
    access = db.get_project_member_role(project_id, current_user.id)
    if project['owner_id'] != current_user.id and access not in ['owner', 'editor'] and not current_user.is_admin:
        abort(403)

    data = request.form if request.form else request.json
    capture_plan_id = data.get("capture_plan_id")

    # Verify capture plan exists and user has access
    plan = db.get_capture_plan(capture_plan_id)
    if not plan:
        return jsonify({"status": "error", "message": "Capture plan not found"}), 404

    # Only capture managers can link projects to plans
    if not current_user.is_capture_manager and not current_user.is_admin:
        return jsonify({"status": "error", "message": "Only capture managers can link projects"}), 403

    db.link_project_to_capture_plan(project_id, capture_plan_id)

    db.write_audit_log("PROJECT_LINKED_TO_CAPTURE",
                     username=current_user.username, user_id=current_user.id,
                     detail=f"Linked project '{project['name']}' to capture plan '{plan['capture_name']}'",
                     ip_address=request.remote_addr)

    flash("Project linked to capture plan.", "success")

    if request.is_json:
        return jsonify({"status": "success"})
    else:
        return redirect(url_for("project_detail", project_id=project_id))


# ── Solicitations ──────────────────────────────────────────────────────────────

@app.route("/solicitations")
def solicitations():
    agency   = request.args.get("agency", "")
    phase    = request.args.get("phase", "")
    program  = request.args.get("program", "")
    source   = request.args.get("source", "")
    status   = request.args.get("status", "")
    keyword  = request.args.get("keyword", "")
    favorited = request.args.get("favorited", "")
    page     = int(request.args.get("page", 1))
    per_page = 50
    offset   = (page - 1) * per_page

    rows = db.get_solicitations(
        agency=agency or None,
        phase=phase or None,
        program=program or None,
        source=source or None,
        status=status or None,
        keyword=keyword or None,
        favorited=True if favorited == "1" else None,
        limit=per_page,
        offset=offset,
    )
    filters = {
        "agencies":  db.get_distinct("solicitations", "agency"),
        "phases":    db.get_distinct("solicitations", "phase"),
        "programs":  db.get_distinct("solicitations", "program"),
        "sources":   db.get_distinct("solicitations", "source"),
        "statuses":  db.get_distinct("solicitations", "status"),
    }
    return render_template("solicitations.html",
                           rows=rows, filters=filters, page=page, per_page=per_page,
                           agency=agency, phase=phase, program=program,
                           source=source, status=status, keyword=keyword,
                           favorited=favorited)


@app.route("/solicitations/<int:sol_id>")
def solicitation_detail(sol_id: int):
    sol = db.get_solicitation(sol_id)
    if not sol:
        flash("Solicitation not found.", "warning")
        return redirect(url_for("solicitations"))
    topics = db.get_topics(keyword=sol.get("solicitation_number") or sol.get("title", "")[:30])
    return render_template("solicitation_detail.html", sol=sol, topics=topics)


@app.route("/solicitations/<int:sol_id>/favorite", methods=["POST"])
def toggle_favorite(sol_id: int):
    db.toggle_favorite(sol_id)
    return redirect(request.referrer or url_for("solicitations"))


@app.route("/solicitations/<int:sol_id>/score", methods=["POST"])
def set_score(sol_id: int):
    score = float(request.form.get("score", 0))
    db.set_score(sol_id, score)
    return redirect(request.referrer or url_for("solicitation_detail", sol_id=sol_id))


@app.route("/solicitations/<int:sol_id>/notes", methods=["POST"])
def set_notes(sol_id: int):
    notes = request.form.get("notes", "")
    db.set_notes(sol_id, notes)
    return redirect(request.referrer or url_for("solicitation_detail", sol_id=sol_id))


# ── Awards ─────────────────────────────────────────────────────────────────────

@app.route("/awards")
def awards():
    agency  = request.args.get("agency", "")
    phase   = request.args.get("phase", "")
    program = request.args.get("program", "")
    year    = request.args.get("year", "")
    source  = request.args.get("source", "")
    keyword = request.args.get("keyword", "")
    page    = int(request.args.get("page", 1))
    per_page = 50
    offset  = (page - 1) * per_page

    rows = db.get_awards(
        agency=agency or None,
        phase=phase or None,
        program=program or None,
        year=int(year) if year else None,
        source=source or None,
        keyword=keyword or None,
        limit=per_page,
        offset=offset,
    )
    filters = {
        "agencies": db.get_distinct("awards", "agency"),
        "phases":   db.get_distinct("awards", "phase"),
        "programs": db.get_distinct("awards", "program"),
        "years":    db.get_distinct("awards", "award_year"),
        "sources":  db.get_distinct("awards", "source"),
    }
    return render_template("awards.html",
                           rows=rows, filters=filters, page=page, per_page=per_page,
                           agency=agency, phase=phase, program=program,
                           year=year, source=source, keyword=keyword)


@app.route("/awards/<int:award_id>")
def award_detail(award_id: int):
    award = db.get_award(award_id)
    if not award:
        flash("Award not found.", "warning")
        return redirect(url_for("awards"))
    return render_template("award_detail.html", award=award)


# ── Topics ─────────────────────────────────────────────────────────────────────

@app.route("/topics/<int:topic_id>")
def topic_detail(topic_id: int):
    import json
    topic = db.get_topic(topic_id)
    if not topic:
        flash("Topic not found.", "warning")
        return redirect(url_for("topics"))
    try:
        topic["ref_docs_parsed"] = json.loads(topic.get("ref_docs") or "[]")
    except Exception:
        topic["ref_docs_parsed"] = []
    return render_template("topic_detail.html", topic=topic)


@app.route("/topics/<int:topic_id>/favorite", methods=["POST"])
@login_required
def toggle_topic_favorite(topic_id: int):
    db.toggle_user_topic_favorite(current_user.id, topic_id)
    return redirect(request.referrer or url_for("topic_detail", topic_id=topic_id))


@app.route("/topics/<int:topic_id>/score", methods=["POST"])
@login_required
def set_topic_score(topic_id: int):
    score = float(request.form.get("score", 0))
    db.set_topic_score(topic_id, score)
    return redirect(request.referrer or url_for("topic_detail", topic_id=topic_id))


@app.route("/topics/<int:topic_id>/notes", methods=["POST"])
@login_required
def set_topic_notes(topic_id: int):
    notes = request.form.get("notes", "")
    db.set_topic_notes(topic_id, notes)
    return redirect(request.referrer or url_for("topic_detail", topic_id=topic_id))


@app.route("/topics/<int:topic_id>/status", methods=["POST"])
@login_required
def set_topic_status(topic_id: int):
    status = request.form.get("status", "")
    db.set_user_topic_status(current_user.id, topic_id, status)
    return redirect(request.referrer or url_for("topics"))


@app.route("/topics/<int:topic_id>/export/pdf")
def export_topic_pdf(topic_id: int):
    """Generate and stream a PDF of the topic detail."""
    import json as _json
    from io import BytesIO
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, HRFlowable)
    from reportlab.lib.enums import TA_LEFT, TA_CENTER

    topic = db.get_topic(topic_id)
    if not topic:
        flash("Topic not found.", "warning")
        return redirect(url_for("topics"))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=0.85*inch, rightMargin=0.85*inch,
                            topMargin=0.85*inch, bottomMargin=0.85*inch)

    styles = getSampleStyleSheet()
    navy   = colors.HexColor("#003087")
    gold   = colors.HexColor("#c8972b")

    title_style = ParagraphStyle("TopicTitle",
                                 fontName="Helvetica-Bold", fontSize=14,
                                 textColor=navy, leading=18, spaceAfter=4)
    meta_style  = ParagraphStyle("Meta",
                                 fontName="Helvetica", fontSize=9,
                                 textColor=colors.HexColor("#6c757d"), spaceAfter=8)
    section_style = ParagraphStyle("Section",
                                   fontName="Helvetica-Bold", fontSize=11,
                                   textColor=navy, spaceBefore=14, spaceAfter=4)
    body_style  = ParagraphStyle("Body",
                                 fontName="Helvetica", fontSize=9,
                                 leading=13, spaceAfter=6)

    def _safe(text):
        """Escape XML special chars for ReportLab Paragraph."""
        if not text:
            return ""
        return (str(text)
                .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"))

    story = []

    # Header bar
    story.append(Paragraph(_safe(topic.get("title", "Untitled Topic")), title_style))
    meta_parts = []
    if topic.get("topic_number"):
        meta_parts.append(f"#{topic['topic_number']}")
    if topic.get("agency"):
        meta_parts.append(topic["agency"])
    if topic.get("branch"):
        meta_parts.append(topic["branch"])
    if topic.get("phase"):
        meta_parts.append(f"Phase {topic['phase']}")
    if topic.get("source"):
        meta_parts.append(topic["source"])
    story.append(Paragraph("  ·  ".join(meta_parts), meta_style))
    story.append(HRFlowable(width="100%", thickness=1, color=gold, spaceAfter=10))

    # Quick-facts table
    facts = []
    if topic.get("close_date"):
        facts.append(("Close Date", topic["close_date"]))
    if topic.get("open_date"):
        facts.append(("Open Date", topic["open_date"]))
    if topic.get("solicitation_year"):
        facts.append(("Sol. Year", topic["solicitation_year"]))
    if topic.get("solicitation_status"):
        facts.append(("Status", topic["solicitation_status"]))
    if topic.get("tech_areas"):
        facts.append(("Tech Area", topic["tech_areas"]))
    if topic.get("itar"):
        facts.append(("ITAR", "Yes"))
    if topic.get("cmmc_level"):
        facts.append(("CMMC Level", topic["cmmc_level"]))
    if topic.get("tech_contact"):
        facts.append(("TPOC", topic["tech_contact"]))
    if topic.get("url"):
        facts.append(("Source URL", topic["url"]))

    if facts:
        tdata = [[Paragraph(f"<b>{_safe(k)}</b>", body_style),
                  Paragraph(_safe(v), body_style)]
                 for k, v in facts]
        tbl = Table(tdata, colWidths=[1.4*inch, 5.3*inch])
        tbl.setStyle(TableStyle([
            ("VALIGN",       (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS",(0, 0), (-1, -1),
             [colors.HexColor("#f8f9fa"), colors.white]),
            ("GRID",         (0, 0), (-1, -1), 0.5, colors.HexColor("#dee2e6")),
            ("LEFTPADDING",  (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING",   (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",(0, 0), (-1, -1), 4),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 10))

    # Sections
    sections = [
        ("Description",     topic.get("description")),
        ("Objective",       topic.get("objective")),
        ("Phase I",         topic.get("phase1_desc")),
        ("Phase II",        topic.get("phase2_desc")),
        ("Phase III",       topic.get("phase3_desc")),
    ]
    for heading, content in sections:
        if content and content.strip():
            story.append(Paragraph(heading, section_style))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.HexColor("#dee2e6"), spaceAfter=4))
            story.append(Paragraph(_safe(content.strip()), body_style))

    # Keywords
    if topic.get("keywords"):
        story.append(Paragraph("Keywords", section_style))
        story.append(Paragraph(_safe(topic["keywords"]), body_style))

    # Notes (if any)
    if topic.get("notes"):
        story.append(Paragraph("My Notes", section_style))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=gold, spaceAfter=4))
        story.append(Paragraph(_safe(topic["notes"]), body_style))

    # Footer note
    story.append(Spacer(1, 16))
    story.append(Paragraph(
        f"Exported from SBIR Pipeline · {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC",
        ParagraphStyle("Footer", fontName="Helvetica", fontSize=8,
                       textColor=colors.HexColor("#adb5bd"))))

    doc.build(story)
    buf.seek(0)

    safe_num = (topic.get("topic_number") or str(topic_id)).replace("/", "-")
    filename = f"SBIR_Topic_{safe_num}.pdf"
    return Response(buf.read(), mimetype="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@app.route("/topics/<int:topic_id>/export/docx")
def export_topic_docx(topic_id: int):
    """Generate and stream a DOCX of the topic detail."""
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    topic = db.get_topic(topic_id)
    if not topic:
        flash("Topic not found.", "warning")
        return redirect(url_for("topics"))

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)

    NAVY = RGBColor(0x00, 0x30, 0x87)
    GOLD = RGBColor(0xC8, 0x97, 0x2B)
    GREY = RGBColor(0x6c, 0x75, 0x7d)

    def _heading(text: str, level: int = 1, color=NAVY):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
        return p

    def _para(text: str, size: int = 10, color=None, italic=False, bold=False):
        p = doc.add_paragraph()
        run = p.add_run(text or "")
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        run.italic = italic
        run.bold   = bold
        return p

    def _add_section(title: str, content: str):
        if not content or not content.strip():
            return
        _heading(title, level=2)
        _para(content.strip(), size=10)

    # Title
    _heading(topic.get("title") or "Untitled Topic", level=1)

    # Meta line
    meta_parts = []
    if topic.get("topic_number"):
        meta_parts.append(f"#{topic['topic_number']}")
    if topic.get("agency"):
        meta_parts.append(topic["agency"])
    if topic.get("branch"):
        meta_parts.append(topic["branch"])
    if topic.get("phase"):
        meta_parts.append(f"Phase {topic['phase']}")
    if topic.get("source"):
        meta_parts.append(topic["source"])
    _para("  ·  ".join(meta_parts), size=9, color=GREY, italic=True)

    doc.add_paragraph()  # spacer

    # Key facts table
    facts = []
    for label, key in [
        ("Close Date",   "close_date"),
        ("Open Date",    "open_date"),
        ("Sol. Year",    "solicitation_year"),
        ("Status",       "solicitation_status"),
        ("Tech Area",    "tech_areas"),
        ("TPOC",         "tech_contact"),
        ("Source URL",   "url"),
    ]:
        val = topic.get(key)
        if val:
            facts.append((label, val))
    if topic.get("itar"):
        facts.append(("ITAR", "Yes"))
    if topic.get("cmmc_level"):
        facts.append(("CMMC Level", topic["cmmc_level"]))

    if facts:
        tbl = doc.add_table(rows=len(facts), cols=2)
        tbl.style = "Table Grid"
        for i, (label, value) in enumerate(facts):
            cell_l = tbl.rows[i].cells[0]
            cell_r = tbl.rows[i].cells[1]
            run_l = cell_l.paragraphs[0].add_run(label)
            run_l.bold = True
            run_l.font.size = Pt(9)
            run_l.font.color.rgb = NAVY
            run_r = cell_r.paragraphs[0].add_run(str(value))
            run_r.font.size = Pt(9)
        doc.add_paragraph()

    # Content sections
    _add_section("Description",  topic.get("description"))
    _add_section("Objective",    topic.get("objective"))
    _add_section("Phase I",      topic.get("phase1_desc"))
    _add_section("Phase II",     topic.get("phase2_desc"))
    _add_section("Phase III",    topic.get("phase3_desc"))

    if topic.get("keywords"):
        _heading("Keywords", level=2)
        _para(topic["keywords"], size=10)

    if topic.get("notes"):
        _heading("My Notes", level=2, color=GOLD)
        _para(topic["notes"], size=10)

    # Footer
    doc.add_paragraph()
    _para(
        f"Exported from SBIR Pipeline · {datetime.utcnow().strftime('%Y-%m-%d %H:%M')} UTC",
        size=8, color=GREY, italic=True,
    )

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_num = (topic.get("topic_number") or str(topic_id)).replace("/", "-")
    filename = f"SBIR_Topic_{safe_num}.docx"
    return Response(
        buf.read(),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.route("/topics")
@login_required
def topics():
    agency       = request.args.get("agency", "")
    phase        = request.args.get("phase", "")
    source       = request.args.get("source", "")
    keyword      = request.args.get("keyword", "")
    favorited    = request.args.get("favorited", "")
    topic_status = request.args.get("topic_status", "")
    page         = int(request.args.get("page", 1))
    per_page     = 50
    offset       = (page - 1) * per_page

    rows = db.get_topics(
        agency=agency or None,
        phase=phase or None,
        source=source or None,
        keyword=keyword or None,
        favorited=True if favorited == "1" else None,
        topic_status=topic_status if topic_status in ("nominated", "passed") else None,
        limit=per_page,
        offset=offset,
        user_id=current_user.id,
    )
    filters = {
        "agencies": db.get_distinct("topics", "agency"),
        "phases":   db.get_distinct("topics", "phase"),
        "sources":  db.get_distinct("topics", "source"),
    }
    return render_template("topics.html",
                           rows=rows, filters=filters, page=page, per_page=per_page,
                           agency=agency, phase=phase, source=source,
                           keyword=keyword, favorited=favorited,
                           topic_status=topic_status)


# ── Topics CSV Export ──────────────────────────────────────────────────────────

@app.route("/topics/export.csv")
@login_required
def export_topics_csv():
    """Export all topics matching current filters as a CSV download."""
    agency       = request.args.get("agency", "")
    phase        = request.args.get("phase", "")
    source       = request.args.get("source", "")
    keyword      = request.args.get("keyword", "")
    favorited    = request.args.get("favorited", "")
    topic_status = request.args.get("topic_status", "")

    rows = db.get_topics(
        agency=agency or None,
        phase=phase or None,
        source=source or None,
        keyword=keyword or None,
        favorited=True if favorited == "1" else None,
        topic_status=topic_status if topic_status in ("nominated", "passed") else None,
        limit=10000,
        offset=0,
        user_id=current_user.id,
    )

    output = io.StringIO()
    writer = csv.writer(output)

    # Header row
    writer.writerow([
        "Topic #", "Title", "Agency", "Branch", "Phase", "Source",
        "Favorited", "Status"
    ])

    for row in rows:
        status_label = ""
        if row["topic_status"] == "nominated":
            status_label = "Interested"
        elif row["topic_status"] == "passed":
            status_label = "Not Interested"

        writer.writerow([
            row["topic_number"] or "",
            row["title"] or "",
            row["agency"] or "",
            row["branch"] or "",
            row["phase"] or "",
            row["source"] or "",
            "Yes" if row["favorited"] else "No",
            status_label,
        ])

    output.seek(0)
    return Response(
        output.getvalue(),
        mimetype="text/csv",
        headers={"Content-Disposition": "attachment; filename=captureiq_topics.csv"},
    )


# ── Analytics ─────────────────────────────────────────────────────────────────

@app.route("/analytics")
@login_required
def analytics():
    data = db.get_analytics_data()
    return render_template("analytics.html", data=data)


# ── Search ─────────────────────────────────────────────────────────────────────

@app.route("/search")
@login_required
def search():
    keyword = request.args.get("q", "").strip()
    results = None
    if keyword:
        results = db.full_search(keyword)
    return render_template("search.html", keyword=keyword, results=results)


# ── Ingestion ──────────────────────────────────────────────────────────────────

@app.route("/ingest/jobs")
@login_required
def ingest_jobs():
    jobs = db.get_ingest_jobs()
    sources = ["sbir.gov", "Navy", "DoD SBIR/STTR"]
    return render_template("ingest_jobs.html", jobs=jobs, sources=sources)


@app.route("/ingest/jobs/create", methods=["POST"])
@login_required
def create_ingest_job():
    source = request.form.get("source", "").strip()
    schedule_type = request.form.get("schedule_type", "daily")
    run_time = request.form.get("run_time", "02:00")
    run_days = request.form.get("run_days", "daily")
    if source:
        db.create_ingest_job(source, schedule_type, run_time, run_days, current_user.id)
        flash(f"Automated job created for {source}.", "success")
    return redirect(url_for("ingest_jobs"))


@app.route("/ingest/jobs/<int:job_id>/edit", methods=["POST"])
@login_required
def edit_ingest_job(job_id):
    db.update_ingest_job(job_id,
        source=request.form.get("source"),
        schedule_type=request.form.get("schedule_type"),
        run_time=request.form.get("run_time"),
        run_days=request.form.get("run_days"))
    flash("Job updated.", "success")
    return redirect(url_for("ingest_jobs"))


@app.route("/ingest/jobs/<int:job_id>/toggle", methods=["POST"])
@login_required
def toggle_ingest_job(job_id):
    with db.get_db() as conn:
        row = conn.execute("SELECT is_active FROM scheduled_ingest_jobs WHERE id=?", (job_id,)).fetchone()
    if row:
        db.update_ingest_job(job_id, is_active=not row['is_active'])
    return redirect(url_for("ingest_jobs"))


@app.route("/ingest/jobs/<int:job_id>/delete", methods=["POST"])
@login_required
def delete_ingest_job(job_id):
    db.delete_ingest_job(job_id)
    flash("Job deleted.", "info")
    return redirect(url_for("ingest_jobs"))


@app.route("/ingest/jobs/<int:job_id>/run", methods=["POST"])
@login_required
def run_ingest_job_now(job_id):
    """Immediately trigger an ingest job in a background thread."""
    jobs = db.get_ingest_jobs()
    job = next((j for j in jobs if j["id"] == job_id), None)
    if job:
        threading.Thread(
            target=_run_ingest_job,
            args=(job["source"], job["id"]),
            daemon=True
        ).start()
        flash(f"Job '{job['source']}' started — check back shortly for status.", "info")
    else:
        flash("Job not found.", "danger")
    return redirect(url_for("ingest_jobs"))


@app.route("/ingest")
@login_required
def ingest_page():
    recent_logs = []
    with db.get_db() as conn:
        rows = conn.execute("""
            SELECT id, source, records_added, records_updated, errors, started_at, finished_at
            FROM ingest_log ORDER BY id DESC LIMIT 20
        """).fetchall()
        recent_logs = [dict(r) for r in rows]
    return render_template("ingest.html", jobs=_jobs, recent_logs=recent_logs)


@app.route("/ingest/sbir-gov", methods=["POST"])
@login_required
def ingest_sbir_gov():
    source_type = request.form.get("source_type", "solicitations")
    agency  = request.form.get("agency", "")
    phase   = request.form.get("phase", "")
    year    = request.form.get("year", "")
    keyword = request.form.get("keyword", "")
    max_rec = int(request.form.get("max_records", 100))

    job_id = f"sbir_gov_{source_type}_{datetime.utcnow().strftime('%H%M%S')}"
    _jobs[job_id] = {"status": "running", "source": f"sbir.gov/{source_type}", "started": datetime.utcnow().isoformat()}

    def run():
        from ingestors import sbir_gov
        try:
            if source_type == "awards":
                result = sbir_gov.ingest_awards(agency=agency, phase=phase, year=year,
                                                keyword=keyword, max_records=max_rec)
            else:
                result = sbir_gov.ingest_solicitations(agency=agency, phase=phase, year=year,
                                                       keyword=keyword, max_records=max_rec)
            _jobs[job_id].update({
                "status": "done",
                "added": result["added"],
                "updated": result["updated"],
                "errors": result["errors"][:3],
                "finished": datetime.utcnow().isoformat(),
            })
        except Exception as e:
            _jobs[job_id].update({"status": "error", "error": str(e)})

    threading.Thread(target=run, daemon=True).start()
    flash(f"Ingestion started (job: {job_id}). Refresh the page to see progress.", "info")
    return redirect(url_for("ingest_page"))


@app.route("/ingest/sbir-topics", methods=["POST"])
@login_required
def ingest_sbir_topics():
    agency    = request.form.get("agency", "").strip()
    phase     = request.form.get("phase", "").strip()
    year      = request.form.get("year", "").strip()
    keyword   = request.form.get("keyword", "").strip()
    status    = request.form.get("status", "open")
    max_rec   = int(request.form.get("max_records", 100))
    fetch_det = request.form.get("fetch_details", "1") == "1"

    job_id = f"sbir_topics_{datetime.utcnow().strftime('%H%M%S')}"
    _jobs[job_id] = {
        "status": "running",
        "source": "sbir.gov/topics",
        "started": datetime.utcnow().isoformat(),
    }
    _ingest_user = (current_user.username, current_user.id)
    db.write_audit_log("INGEST_STARTED",
                       username=_ingest_user[0], user_id=_ingest_user[1],
                       detail=f"Source: sbir.gov/topics, max={max_rec}",
                       ip_address=request.remote_addr)

    def run():
        from ingestors import sbir_gov_topics
        try:
            result = sbir_gov_topics.ingest(
                agency=agency, phase=phase, year=year,
                keyword=keyword, status=status,
                max_records=max_rec, fetch_details=fetch_det,
            )
            _jobs[job_id].update({
                "status": "done",
                "added": result["added"],
                "updated": result["updated"],
                "errors": result["errors"][:3],
                "finished": datetime.utcnow().isoformat(),
            })
            db.write_audit_log("INGEST_COMPLETED",
                               username=_ingest_user[0], user_id=_ingest_user[1],
                               detail=f"sbir.gov/topics: added={result['added']}, updated={result['updated']}")
        except Exception as e:
            _jobs[job_id].update({"status": "error", "error": str(e)})
            db.write_audit_log("INGEST_FAILED",
                               username=_ingest_user[0], user_id=_ingest_user[1],
                               detail=f"sbir.gov/topics error: {e}")

    threading.Thread(target=run, daemon=True).start()
    flash(f"SBIR.gov topics ingestion started (job: {job_id}).", "info")
    return redirect(url_for("ingest_page"))


@app.route("/ingest/navy", methods=["POST"])
@login_required
def ingest_navy():
    topics_url = request.form.get("topics_url", "").strip() or "https://www.navysbir.com/topics26_1.htm"
    max_topics = int(request.form.get("max_topics", 100))

    job_id = f"navy_sbir_{datetime.utcnow().strftime('%H%M%S')}"
    _jobs[job_id] = {"status": "running", "source": "navysbir.com", "started": datetime.utcnow().isoformat()}

    def run():
        from ingestors import navy_sbir
        try:
            result = navy_sbir.ingest(topics_url=topics_url, max_topics=max_topics)
            _jobs[job_id].update({
                "status": "done",
                "added": result["added"],
                "updated": result["updated"],
                "errors": result["errors"][:3],
                "finished": datetime.utcnow().isoformat(),
            })
        except Exception as e:
            _jobs[job_id].update({"status": "error", "error": str(e)})

    threading.Thread(target=run, daemon=True).start()
    flash(f"Navy SBIR ingestion started (job: {job_id}).", "info")
    return redirect(url_for("ingest_page"))


@app.route("/ingest/dod", methods=["POST"])
@login_required
def ingest_dod():
    baa     = request.form.get("baa", "DOD_SBIR_2026_P1_CBZ").strip()
    keyword = request.form.get("keyword", "").strip()
    max_rec = int(request.form.get("max_records", 200))

    job_id = f"dod_{baa}_{datetime.utcnow().strftime('%H%M%S')}"
    _jobs[job_id] = {
        "status": "running",
        "source": f"dod_sbirsttr/{baa}",
        "started": datetime.utcnow().isoformat(),
    }

    def run():
        from ingestors import dod_sbirsttr
        try:
            result = dod_sbirsttr.ingest(baa=baa, keyword=keyword, max_records=max_rec)
            _jobs[job_id].update({
                "status": "done",
                "added": result["added"],
                "updated": result["updated"],
                "errors": result["errors"][:5],
                "endpoint_used": result.get("endpoint_used"),
                "finished": datetime.utcnow().isoformat(),
            })
        except Exception as e:
            _jobs[job_id].update({"status": "error", "error": str(e)})

    threading.Thread(target=run, daemon=True).start()
    flash(f"DoD SBIR/STTR ingestion started for BAA={baa} (job: {job_id}).", "info")
    return redirect(url_for("ingest_page"))


@app.route("/api/dod/probe")
def api_dod_probe():
    """Diagnostic: test which DoD API endpoints are reachable."""
    from ingestors import dod_sbirsttr
    baa = request.args.get("baa", "DOD_SBIR_2026_P1_CBZ")
    results = dod_sbirsttr.probe_endpoints()
    return jsonify({"baa": baa, "endpoints": results})


@app.route("/api/dod/baas")
def api_dod_baas():
    """Return the list of available BAA identifiers from the DoD portal."""
    from ingestors import dod_sbirsttr
    baas = dod_sbirsttr.get_available_baas()
    return jsonify(baas)


# ── SBIR Capture — Projects ────────────────────────────────────────────────────

@app.route("/projects")
@login_required
def projects():
    stage   = request.args.get("stage", "")
    keyword = request.args.get("keyword", "")
    rows = db.get_projects(
        stage=stage or None,
        keyword=keyword or None,
    )
    capture_stats = db.get_capture_stats()
    return render_template("projects.html",
                           rows=rows,
                           stages=db.STAGES,
                           capture_stats=capture_stats,
                           stage=stage,
                           keyword=keyword)


@app.route("/projects/new", methods=["POST"])
@login_required
def create_project():
    topic_id = request.form.get("topic_id") or None
    if topic_id:
        try:
            topic_id = int(topic_id)
        except ValueError:
            topic_id = None

    name = request.form.get("name", "").strip()
    if not name:
        flash("Project name is required.", "warning")
        return redirect(url_for("projects"))

    project_id = db.create_project({
        "topic_id":       topic_id,
        "name":           name,
        "description":    request.form.get("description", "").strip(),
        "stage":          request.form.get("stage", "Identified"),
        "lead":           request.form.get("lead", "").strip(),
        "due_date":       request.form.get("due_date", "").strip() or None,
        "checklist_type": request.form.get("checklist_type", "dod"),
        "source":         request.form.get("source", "").strip(),
    })

    # Set the current user as the project owner
    db.set_project_owner(project_id, current_user.id)

    # Auto-add the creator as an active team member (lead role) so they
    # can access the team page and appear in team listings immediately
    db.add_team_member(project_id, current_user.id, role='lead',
                       added_by_user_id=current_user.id)

    # Log the action
    db.log_activity(project_id, "created", f"Project created by {current_user.username}")

    flash(f"Project '{name}' created.", "success")
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>")
@require_project_access(required_role='viewer')
def project_detail(project_id: int):
    from flask import g
    project = db.get_project(project_id)
    if not project:
        flash("Project not found.", "warning")
        return redirect(url_for("projects"))

    checklist = db.get_checklist(project_id)
    files     = db.get_project_files(project_id)
    activity  = db.get_activity_log(project_id)
    members   = db.get_project_members(project_id)
    comments  = db.get_project_comments(project_id)
    shared_documents = db.get_shared_documents(project_id)
    team_members = db.get_project_team_members(project_id)
    pending_invitations = db.get_project_invitations(project_id, status='pending')

    # Get user's pending invitations (for this project and others)
    user_pending_invs = db.get_pending_invitations(current_user.id)
    user_pending_invitations = {inv['project_id']: inv['id'] for inv in user_pending_invs}

    # Group checklist by category
    checklist_groups = {}
    for item in checklist:
        cat = item["category"] or "General"
        checklist_groups.setdefault(cat, []).append(item)

    gantt_items  = db.get_checklist_items_for_gantt(project_id)
    all_users    = db.get_all_users()

    return render_template("project_detail.html",
                           project=project,
                           checklist_groups=checklist_groups,
                           gantt_items=gantt_items,
                           all_users=all_users,
                           files=files,
                           activity=activity,
                           stages=db.STAGES,
                           members=members,
                           comments=comments,
                           shared_documents=shared_documents,
                           user_role=g.project_user_role,
                           team_members=team_members,
                           pending_invitations=pending_invitations,
                           user_pending_invitations=user_pending_invitations)


@app.route("/projects/<int:project_id>/edit", methods=["POST"])
@login_required
def edit_project(project_id: int):
    db.update_project(project_id, {
        "name":        request.form.get("name", "").strip(),
        "description": request.form.get("description", "").strip(),
        "lead":        request.form.get("lead", "").strip(),
        "due_date":    request.form.get("due_date", "").strip() or None,
        "notes":       request.form.get("notes", "").strip(),
    })
    flash("Project updated.", "success")
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/stage", methods=["POST"])
@login_required
def set_project_stage(project_id: int):
    stage = request.form.get("stage", "Identified")
    db.set_project_stage(project_id, stage)
    flash(f"Stage updated to '{stage}'.", "success")
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/delete", methods=["POST"])
@require_project_access(required_role='owner')
def delete_project(project_id: int):
    project = db.get_project(project_id)
    if project:
        # Delete uploaded files from disk
        files = db.get_project_files(project_id)
        for f in files:
            try:
                if os.path.exists(f.get("local_path")):
                    os.remove(f["local_path"])
            except Exception:
                pass
        db.delete_project(project_id)
        db.log_activity(project_id, "deleted", f"Project deleted by {current_user.username}")
        flash(f"Project '{project['name']}' deleted.", "info")
    return redirect(url_for("projects"))


# ── Project Sharing & Collaboration ────────────────────────────────────────────

@app.route("/projects/<int:project_id>/share", methods=["GET", "POST"])
@require_project_access(required_role='owner')
def share_project(project_id: int):
    """Manage project team members and permissions."""
    project = db.get_project(project_id)
    members = db.get_project_members(project_id)
    all_users = db.get_all_users()

    if request.method == "POST":
        user_id = request.form.get("user_id")
        role = request.form.get("role", "viewer")

        if user_id and role in ['viewer', 'editor']:
            try:
                user_id = int(user_id)
                db.add_project_member(project_id, user_id, role, current_user.id)
                username = next((u['username'] for u in all_users if u['id'] == user_id), "User")
                db.log_activity(project_id, "shared", f"{username} added with {role} role by {current_user.username}")

                # Create notification
                db.create_notification(user_id, 'share', project_id, current_user.id,
                                     f"{current_user.username} shared project '{project['name']}' with you")

                flash(f"User added to project with {role} access.", "success")
            except Exception as e:
                flash(f"Error adding user: {str(e)}", "danger")

    return render_template("project_share.html",
                          project=project,
                          members=members,
                          all_users=[u for u in all_users if u['id'] != project['owner_id']])


@app.route("/projects/<int:project_id>/members/<int:user_id>/remove", methods=["POST"])
@require_project_access(required_role='owner')
def remove_project_member(project_id: int, user_id: int):
    """Remove a user from a project."""
    db.remove_project_member(project_id, user_id)
    project = db.get_project(project_id)
    user = db.get_user_by_id(user_id)
    if user:
        db.log_activity(project_id, "shared", f"{user['username']} removed from project by {current_user.username}")
    flash("Team member removed.", "success")
    return redirect(url_for("share_project", project_id=project_id))


@app.route("/projects/<int:project_id>/members/<int:user_id>/role", methods=["POST"])
@require_project_access(required_role='owner')
def update_member_role(project_id: int, user_id: int):
    """Update a team member's role in the project."""
    role = request.form.get("role", "viewer")
    if role in ['viewer', 'editor']:
        db.update_project_member_role(project_id, user_id, role)
        user = db.get_user_by_id(user_id)
        if user:
            db.log_activity(project_id, "shared", f"{user['username']} role updated to {role} by {current_user.username}")
        flash(f"Role updated to {role}.", "success")
    return redirect(url_for("share_project", project_id=project_id))


# ── Project Comments ───────────────────────────────────────────────────────────

@app.route("/projects/<int:project_id>/comments/add", methods=["POST"])
@require_project_access(required_role='viewer')
def add_comment(project_id: int):
    """Add a comment to a project."""
    comment_text = request.form.get("comment_text", "").strip()
    file_id = request.form.get("file_id") or None
    if file_id:
        try:
            file_id = int(file_id)
        except (ValueError, TypeError):
            file_id = None

    if not comment_text:
        flash("Comment cannot be empty.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    comment_id = db.add_project_comment(project_id, current_user.id, comment_text, file_id)
    project = db.get_project(project_id)

    # Log and notify
    db.log_activity(project_id, "note", f"Comment added by {current_user.username}")

    # Notify other project members
    members = db.get_project_members(project_id)
    for member in members:
        if member['user_id'] != current_user.id:
            db.create_notification(member['user_id'], 'comment', project_id, current_user.id,
                                 f"{current_user.username} commented on '{project['name']}'")

    # Also notify project owner if current user isn't owner
    if project.get('owner_id') and project['owner_id'] != current_user.id:
        db.create_notification(project['owner_id'], 'comment', project_id, current_user.id,
                             f"{current_user.username} commented on '{project['name']}'")

    flash("Comment added.", "success")
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/comments/<int:comment_id>/delete", methods=["POST"])
@login_required
def delete_comment(project_id: int, comment_id: int):
    """Delete a comment (owner or author only)."""
    # Check if user is comment author or project owner
    comment = db.get_project_comment(comment_id)
    project = db.get_project(project_id)

    if not comment or comment.get('project_id') != project_id:
        flash("Comment not found.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    if comment.get('user_id') != current_user.id and project.get('owner_id') != current_user.id:
        flash("You can only delete your own comments.", "danger")
        return redirect(url_for("project_detail", project_id=project_id))

    db.delete_comment(comment_id)
    flash("Comment deleted.", "info")
    return redirect(url_for("project_detail", project_id=project_id))


# ── Notifications ──────────────────────────────────────────────────────────────

@app.route("/api/notifications")
@login_required
def get_notifications():
    """Get user's unread notifications (JSON API)."""
    notifications = db.get_user_notifications(current_user.id, unread_only=True)
    return jsonify(notifications)


@app.route("/api/notifications/<int:notif_id>/read", methods=["POST"])
@login_required
def mark_notification_read(notif_id: int):
    """Mark a notification as read."""
    db.mark_notification_read(notif_id)
    return jsonify({"success": True})


@app.route("/api/notifications/mark-all-read", methods=["POST"])
@login_required
def mark_all_read():
    """Mark all notifications as read."""
    db.mark_all_notifications_read(current_user.id)
    return jsonify({"success": True})


@app.route("/notifications")
@login_required
def notifications():
    """View all notifications."""
    notifications = db.get_user_notifications(current_user.id)
    unread_count = db.get_unread_notification_count(current_user.id)
    return render_template("notifications.html",
                          notifications=notifications,
                          unread_count=unread_count)


# ── SBIR Capture — Files ───────────────────────────────────────────────────────

@app.route("/projects/<int:project_id>/files/upload", methods=["POST"])
@login_required
def upload_project_file(project_id: int):
    from integrations import google_drive as gd

    project = db.get_project(project_id)
    if not project:
        flash("Project not found.", "warning")
        return redirect(url_for("projects"))

    if "file" not in request.files:
        flash("No file selected.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    f = request.files["file"]
    if f.filename == "":
        flash("No file selected.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    if not _allowed_file(f.filename):
        flash(f"File type not allowed. Permitted: {', '.join(sorted(ALLOWED_EXTENSIONS))}", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    filename  = secure_filename(f.filename)
    category  = request.form.get("file_category", "general")
    dest      = request.form.get("destination", "local")

    # Always save locally first (needed for Drive upload too)
    proj_dir  = os.path.join(UPLOAD_FOLDER, str(project_id))
    os.makedirs(proj_dir, exist_ok=True)
    save_path = os.path.join(proj_dir, filename)
    base, ext = os.path.splitext(filename)
    counter = 1
    while os.path.exists(save_path):
        filename  = f"{base}_{counter}{ext}"
        save_path = os.path.join(proj_dir, filename)
        counter  += 1
    f.save(save_path)
    size = os.path.getsize(save_path)

    if dest == "gdrive" and gd.is_connected():
        try:
            # Ensure project has a Drive folder
            folder_id = project.get("gdrive_folder_id")
            if not folder_id:
                folder_id = gd.get_or_create_project_folder(
                    project["name"], project_id)
                db.set_project_gdrive_folder(project_id, folder_id)

            gdrive_id, web_link = gd.upload_file(
                folder_id, save_path, filename, f.content_type)

            # Remove the temporary local copy after upload
            try:
                os.remove(save_path)
            except Exception:
                pass

            db.add_project_file(
                project_id=project_id,
                filename=filename,
                local_path=None,
                file_size=size,
                mime_type=f.content_type,
                category=category,
                storage_backend="gdrive",
                gdrive_file_id=gdrive_id,
                gdrive_web_link=web_link,
            )
            flash(f"'{filename}' uploaded to Google Drive.", "success")
        except Exception as e:
            flash(f"Drive upload failed ({e}). File saved locally instead.", "warning")
            db.add_project_file(
                project_id=project_id,
                filename=filename,
                local_path=save_path,
                file_size=size,
                mime_type=f.content_type,
                category=category,
            )
    else:
        db.add_project_file(
            project_id=project_id,
            filename=filename,
            local_path=save_path,
            file_size=size,
            mime_type=f.content_type,
            category=category,
        )
        flash(f"'{filename}' uploaded successfully.", "success")

    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/files/<int:file_id>/download")
@login_required
def download_project_file(project_id: int, file_id: int):
    from integrations import google_drive as gd
    file_rec = db.get_project_file(file_id, project_id)
    if not file_rec:
        flash("File not found.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    if file_rec.get("storage_backend") == "gdrive":
        try:
            buf = gd.download_file(file_rec["gdrive_file_id"])
            return Response(
                buf.read(),
                mimetype=file_rec.get("mime_type") or "application/octet-stream",
                headers={"Content-Disposition":
                         f'attachment; filename="{file_rec["filename"]}"'},
            )
        except Exception as e:
            flash(f"Could not download from Google Drive: {e}", "danger")
            return redirect(url_for("project_detail", project_id=project_id))

    if not file_rec.get("local_path") or not os.path.exists(file_rec["local_path"]):
        flash("Local file not found.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))
    return send_file(file_rec["local_path"], as_attachment=True,
                     download_name=file_rec["filename"])


@app.route("/projects/<int:project_id>/files/<int:file_id>/delete", methods=["POST"])
@login_required
def delete_project_file(project_id: int, file_id: int):
    from integrations import google_drive as gd
    file_rec = db.get_project_file(file_id, project_id)
    if file_rec:
        if file_rec.get("storage_backend") == "gdrive" and file_rec.get("gdrive_file_id"):
            try:
                gd.delete_file(file_rec["gdrive_file_id"])
            except Exception:
                pass  # Best-effort Drive delete
        local_path = db.delete_project_file(file_id, project_id)
        if local_path:
            try:
                if os.path.exists(local_path):
                    os.remove(local_path)
            except Exception:
                pass
        flash("File deleted.", "info")
    return redirect(url_for("project_detail", project_id=project_id))


# ── SBIR Capture — Checklist ───────────────────────────────────────────────────

@app.route("/projects/<int:project_id>/checklist/<int:item_id>/toggle", methods=["POST"])
def toggle_checklist(project_id: int, item_id: int):
    db.toggle_checklist_item(item_id, project_id)
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/checklist/add", methods=["POST"])
@login_required
def add_checklist_item(project_id: int):
    label    = request.form.get("label", "").strip()
    category = request.form.get("category", "Custom").strip() or "Custom"
    if label:
        db.add_checklist_item(project_id, label, category)
        flash("Checklist item added.", "success")
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/checklist/<int:item_id>/delete", methods=["POST"])
@login_required
def delete_checklist_item(project_id: int, item_id: int):
    db.delete_checklist_item(item_id, project_id)
    return redirect(url_for("project_detail", project_id=project_id))


# ── Google Drive Settings ──────────────────────────────────────────────────────

@app.route("/settings/gdrive")
@login_required
def gdrive_settings():
    from integrations import google_drive as gd
    return render_template("gdrive_settings.html",
                           connected=gd.is_connected(),
                           has_creds=gd.has_credentials_file())


@app.route("/settings/gdrive/connect")
@login_required
def gdrive_connect():
    from integrations import google_drive as gd
    if not gd.has_credentials_file():
        flash("credentials.json not found. See setup instructions.", "danger")
        return redirect(url_for("gdrive_settings"))
    redirect_uri = url_for("gdrive_callback", _external=True)
    return redirect(gd.get_auth_url(redirect_uri))


@app.route("/settings/gdrive/callback")
@login_required
def gdrive_callback():
    from integrations import google_drive as gd
    code = request.args.get("code")
    if not code:
        flash("Google authorisation failed — no code returned.", "danger")
        return redirect(url_for("gdrive_settings"))
    try:
        redirect_uri = url_for("gdrive_callback", _external=True)
        gd.exchange_code(code, redirect_uri)
        flash("Google Drive connected successfully!", "success")
    except Exception as e:
        flash(f"Google Drive connection failed: {e}", "danger")
    return redirect(url_for("gdrive_settings"))


@app.route("/settings/gdrive/disconnect", methods=["POST"])
@login_required
def gdrive_disconnect():
    from integrations import google_drive as gd
    gd.revoke()
    flash("Google Drive disconnected.", "info")
    return redirect(url_for("gdrive_settings"))


# ── SharePoint Settings ────────────────────────────────────────────────────────

@app.route("/settings/sharepoint")
@login_required
def sharepoint_settings():
    from integrations import sharepoint as sp
    return render_template("sharepoint_settings.html",
                          connected=sp.is_connected())


@app.route("/settings/sharepoint/disconnect", methods=["POST"])
@login_required
def sharepoint_disconnect():
    from integrations import sharepoint as sp
    import os
    try:
        if os.path.exists(sp.SHAREPOINT_TOKEN_FILE):
            os.remove(sp.SHAREPOINT_TOKEN_FILE)
        flash("SharePoint disconnected.", "info")
    except Exception as e:
        flash(f"Error disconnecting: {e}", "danger")
    return redirect(url_for("sharepoint_settings"))


# ── Shared Documents ──────────────────────────────────────────────────────────

@app.route("/projects/<int:project_id>/documents/link", methods=["POST"])
@require_project_access(required_role='editor')
def link_document(project_id: int):
    """Link an external document (SharePoint/Drive) to a project."""
    doc_type = request.form.get("doc_type", "").strip()
    external_url = request.form.get("external_url", "").strip()
    title = request.form.get("title", "").strip()
    external_id = request.form.get("external_id", "").strip()

    if not doc_type or not external_url:
        flash("Document type and URL are required.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))

    doc_id = db.add_shared_document(project_id, doc_type, external_url,
                                   external_id, title, current_user.id)
    project = db.get_project(project_id)
    db.log_activity(project_id, "file_upload",
                   f"{title or 'Document'} linked from {doc_type.upper()} by {current_user.username}")

    flash(f"Document linked to project.", "success")
    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/documents/<int:doc_id>/remove", methods=["POST"])
@require_project_access(required_role='editor')
def remove_document(project_id: int, doc_id: int):
    """Remove a document link from a project."""
    docs = db.get_shared_documents(project_id)
    doc = next((d for d in docs if d['id'] == doc_id), None)

    if doc:
        db.remove_shared_document(doc_id)
        db.log_activity(project_id, "file_delete",
                       f"Document '{doc['title']}' removed by {current_user.username}")
        flash("Document removed.", "info")
    else:
        flash("Document not found.", "warning")

    return redirect(url_for("project_detail", project_id=project_id))


# ── Project Team Management ────────────────────────────────────────────────────

@app.route("/projects/<int:project_id>/team", methods=["GET"])
@login_required
def project_team(project_id: int):
    """View project team members and manage invitations."""
    project = db.get_project(project_id)
    if not project:
        abort(404)

    # Check access (project members or admins)
    if not (db.is_project_team_member(project_id, current_user.id) or current_user.is_admin):
        abort(403)

    team_members = db.get_project_team_members(project_id)
    pending_invitations = db.get_project_invitations(project_id, status='pending')

    # Check if current user is project lead or admin (can manage team)
    can_manage = (project.get('created_by_user_id') == current_user.id or
                  current_user.is_admin)

    all_users = [u for u in db.get_all_users() if u['is_active']]
    team_ids = {m['user_id'] for m in team_members}
    available_users = [u for u in all_users if u['id'] not in team_ids]

    return render_template("project_team.html",
                         project=project,
                         team_members=team_members,
                         pending_invitations=pending_invitations,
                         can_manage=can_manage,
                         available_users=available_users)


@app.route("/projects/<int:project_id>/team/invite", methods=["POST"])
@login_required
def invite_team_member(project_id: int):
    """Send a team invitation to a user."""
    project = db.get_project(project_id)
    if not project:
        abort(404)

    # Check if user can manage team (project lead or admin)
    if project.get('created_by_user_id') != current_user.id and not current_user.is_admin:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Not authorized'}), 403
        flash("You don't have permission to manage this team.", "danger")
        return redirect(url_for("project_detail", project_id=project_id))

    user_id = request.form.get('user_id') or (request.json.get('user_id') if request.is_json else None)

    if not user_id:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'User ID required'}), 400
        flash("User not specified.", "warning")
        return redirect(url_for("project_team", project_id=project_id))

    user_id = int(user_id)
    user = db.get_user_by_id(user_id)
    if not user:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'User not found'}), 404
        flash("User not found.", "warning")
        return redirect(url_for("project_team", project_id=project_id))

    # Check if already on team
    if db.is_project_team_member(project_id, user_id):
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Already on team'}), 400
        flash(f"{user['username']} is already on this team.", "info")
        return redirect(url_for("project_team", project_id=project_id))

    # Send invitation
    if db.send_team_invitation(project_id, user_id, current_user.id):
        db.write_audit_log("PROJECT_INVITATION_SENT",
                          username=current_user.username, user_id=current_user.id,
                          detail=f"Invited '{user['username']}' to project '{project['name']}'",
                          ip_address=request.remote_addr)

        if request.is_json:
            return jsonify({'status': 'success', 'message': f'Invitation sent to {user["username"]}'})
        flash(f"Invitation sent to {user['username']}.", "success")
    else:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Failed to send invitation'}), 500
        flash("Failed to send invitation.", "danger")

    return redirect(url_for("project_team", project_id=project_id))


@app.route("/projects/<int:project_id>/team/<int:user_id>/remove", methods=["POST"])
@login_required
def remove_team_member(project_id: int, user_id: int):
    """Remove a user from project team."""
    project = db.get_project(project_id)
    if not project:
        abort(404)

    # Check if user can manage team
    if project.get('created_by_user_id') != current_user.id and not current_user.is_admin:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Not authorized'}), 403
        flash("You don't have permission to manage this team.", "danger")
        return redirect(url_for("project_detail", project_id=project_id))

    user = db.get_user_by_id(user_id)
    if not user:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'User not found'}), 404
        abort(404)

    if db.remove_team_member(project_id, user_id):
        db.write_audit_log("PROJECT_MEMBER_REMOVED",
                          username=current_user.username, user_id=current_user.id,
                          detail=f"Removed '{user['username']}' from project '{project['name']}'",
                          ip_address=request.remote_addr)

        if request.is_json:
            return jsonify({'status': 'success', 'message': f'{user["username"]} removed from team'})
        flash(f"{user['username']} removed from team.", "success")
    else:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Failed to remove member'}), 500
        flash("Failed to remove team member.", "danger")

    return redirect(url_for("project_team", project_id=project_id))


@app.route("/projects/<int:project_id>/invite/<int:invitation_id>/accept", methods=["POST"])
@login_required
def accept_invitation(project_id: int, invitation_id: int):
    """Accept a team invitation."""
    project = db.get_project(project_id)
    if not project:
        abort(404)

    if db.accept_team_invitation(project_id, current_user.id):
        db.write_audit_log("PROJECT_INVITATION_ACCEPTED",
                          username=current_user.username, user_id=current_user.id,
                          detail=f"Accepted invitation to project '{project['name']}'",
                          ip_address=request.remote_addr)

        if request.is_json:
            return jsonify({'status': 'success', 'message': 'Invitation accepted'})
        flash("You've joined the project team!", "success")
    else:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Failed to accept invitation'}), 500
        flash("Failed to accept invitation.", "danger")

    return redirect(url_for("project_detail", project_id=project_id))


@app.route("/projects/<int:project_id>/invite/<int:invitation_id>/decline", methods=["POST"])
@login_required
def decline_invitation(project_id: int, invitation_id: int):
    """Decline a team invitation."""
    project = db.get_project(project_id)
    if not project:
        abort(404)

    if db.decline_team_invitation(project_id, current_user.id):
        db.write_audit_log("PROJECT_INVITATION_DECLINED",
                          username=current_user.username, user_id=current_user.id,
                          detail=f"Declined invitation to project '{project['name']}'",
                          ip_address=request.remote_addr)

        if request.is_json:
            return jsonify({'status': 'success', 'message': 'Invitation declined'})
        flash("You've declined the invitation.", "info")
    else:
        if request.is_json:
            return jsonify({'status': 'error', 'message': 'Failed to decline invitation'}), 500
        flash("Failed to decline invitation.", "danger")

    return redirect(url_for("dashboard"))


# ── API endpoints (JSON) ───────────────────────────────────────────────────────

@app.route("/ingest-log/<int:log_id>/delete", methods=["POST"])
@login_required
def delete_ingest_log(log_id: int):
    db.delete_ingest_log(log_id)
    return redirect(url_for("ingest_page"))


@app.route("/api/project/<int:project_id>/available-users")
@login_required
def api_available_users(project_id: int):
    """Get users not on project team (for invitations)."""
    project = db.get_project(project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404

    # Get all users
    all_users = db.get_all_users()

    # Get team members
    team_members = db.get_project_team_members(project_id)
    team_member_ids = {m['user_id'] for m in team_members}

    # Get pending invitations
    pending_invitations = db.get_project_invitations(project_id, status='pending')
    pending_user_ids = {i['invited_user_id'] for i in pending_invitations}

    # Filter out current user, team members, and already invited
    available = [
        {'id': u['id'], 'username': u['username'], 'email': u['email']}
        for u in all_users
        if u['id'] != current_user.id  # Not current user
        and u['id'] not in team_member_ids  # Not already on team
        and u['id'] not in pending_user_ids  # Not already invited
    ]

    return jsonify({'users': available})


@app.route("/api/stats")
def api_stats():
    return jsonify(db.get_stats())


@app.route("/api/jobs")
def api_jobs():
    return jsonify(_jobs)


@app.route("/api/solicitations")
def api_solicitations():
    keyword = request.args.get("q", "")
    agency  = request.args.get("agency", "")
    phase   = request.args.get("phase", "")
    rows = db.get_solicitations(
        keyword=keyword or None,
        agency=agency or None,
        phase=phase or None,
        limit=100,
    )
    return jsonify(rows)


@app.route("/api/awards")
def api_awards():
    keyword = request.args.get("q", "")
    agency  = request.args.get("agency", "")
    rows = db.get_awards(keyword=keyword or None, agency=agency or None, limit=100)
    return jsonify(rows)


# ── Automated Ingest Job Scheduler ────────────────────────────────────────────

def _run_ingest_job(source: str, job_db_id: int):
    """Execute an ingest job for the given source and update last_run / last_status in DB."""
    now = datetime.now().isoformat()  # local time, matches scheduler tick comparison
    status = "error"
    try:
        src = source.lower()
        if "navy" in src:
            from ingestors import navy_sbir
            navy_sbir.ingest(max_topics=200)
        elif "dod" in src:
            from ingestors import dod_sbirsttr
            dod_sbirsttr.ingest(max_records=200)
        elif "sbir.gov" in src or "sbir gov" in src:
            from ingestors import sbir_gov
            sbir_gov.ingest_topics(limit=200)
        status = "success"
        print(f"[Scheduler] Job {job_db_id} ({source}) completed successfully.")
    except Exception as e:
        print(f"[Scheduler] Job {job_db_id} ({source}) failed: {e}")
    finally:
        with db.get_db() as conn:
            conn.execute(
                "UPDATE scheduled_ingest_jobs SET last_run=?, last_status=?, updated_at=? WHERE id=?",
                (now, status, now, job_db_id)
            )


def _scheduler_loop():
    """Background thread: checks every 60s for jobs due to run.
    Uses local machine time so run_time values set in the UI match the user's timezone.
    """
    import time as _time
    from datetime import datetime as _dt
    print("[Scheduler] Automated ingest scheduler started (local time).")
    while True:
        try:
            # Use local time so scheduled times match what the user configured in the UI
            now = _dt.now()
            day_name = now.strftime("%A").lower()   # e.g. "monday"
            hm = now.strftime("%H:%M")               # e.g. "14:30"
            print(f"[Scheduler] Tick — local time: {now.strftime('%Y-%m-%d %H:%M')}, checking jobs...")

            jobs = db.get_ingest_jobs()
            for job in jobs:
                if not job.get("is_active"):
                    continue

                run_time  = (job.get("run_time") or "02:00").strip()
                run_days  = (job.get("run_days") or "daily").strip().lower()
                sched     = (job.get("schedule_type") or "daily").strip().lower()

                print(f"[Scheduler]   Job {job['id']} ({job['source']}): run_time={run_time}, now={hm}, match={run_time == hm}")

                # Check time window (within the current minute)
                if run_time != hm:
                    continue

                # Check day
                if sched == "monthly":
                    if now.day != 1:
                        print(f"[Scheduler]   Job {job['id']}: monthly but not 1st of month, skipping.")
                        continue
                elif run_days == "weekdays":
                    if day_name not in ("monday","tuesday","wednesday","thursday","friday"):
                        print(f"[Scheduler]   Job {job['id']}: weekdays only, today is {day_name}, skipping.")
                        continue
                elif run_days not in ("daily",):
                    # Specific day like "monday", "tuesday" etc.
                    if run_days != day_name:
                        print(f"[Scheduler]   Job {job['id']}: scheduled for {run_days}, today is {day_name}, skipping.")
                        continue
                # "daily" falls through — always fires at the right time

                # Avoid double-firing in the same minute
                last_run = job.get("last_run") or ""
                this_minute = now.strftime("%Y-%m-%dT%H:%M")
                if last_run.startswith(this_minute):
                    print(f"[Scheduler]   Job {job['id']}: already ran this minute, skipping.")
                    continue

                print(f"[Scheduler] >>> Firing job {job['id']} — {job['source']}")
                threading.Thread(
                    target=_run_ingest_job,
                    args=(job["source"], job["id"]),
                    daemon=True
                ).start()

        except Exception as e:
            import traceback
            print(f"[Scheduler] Error in scheduler loop: {e}")
            traceback.print_exc()

        _time.sleep(60)


def start_scheduler():
    """Start the background ingest scheduler thread.

    Guards against Flask debug reloader's double-process issue:
    when debug=True, werkzeug spawns a child process (WERKZEUG_RUN_MAIN=true).
    We only want the scheduler in the child worker, not the parent reloader.
    When debug=False (launcher.py), WERKZEUG_RUN_MAIN is not set, so we start normally.
    """
    import os
    debug_mode = app.debug
    if debug_mode and os.environ.get("WERKZEUG_RUN_MAIN") != "true":
        # Parent reloader process — skip starting scheduler here
        print("[Scheduler] Skipping scheduler start in reloader parent process.")
        return
    t = threading.Thread(target=_scheduler_loop, daemon=True)
    t.start()


# ── Entry point ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    db.init_db()
    start_scheduler()
    print("\n" + "="*60)
    print("  CaptureIQ")
    print("  Open http://127.0.0.1:5000 in your browser")
    print("="*60 + "\n")
    app.run(debug=True, host="127.0.0.1", port=5000)


# ── Proposal Scoring & Ranking ───────────────────────────────────────────────

@app.route("/capture-plans/<int:plan_id>/scoring", methods=["GET"])
@login_required
@require_capture_manager
def scoring_dashboard(plan_id: int):
    """View scoring dashboard for a capture plan."""
    plan = db.get_capture_plan(plan_id)
    if not plan:
        flash("Capture plan not found.", "warning")
        return redirect(url_for("capture_dashboard"))
    
    criteria = db.get_scoring_criteria(plan_id)
    progress = db.get_scoring_progress(plan_id)
    rankings = db.get_capture_plan_rankings(plan_id)
    
    return render_template("scoring_dashboard.html",
                          plan=plan,
                          criteria=criteria,
                          progress=progress,
                          rankings=rankings)


@app.route("/capture-plans/<int:plan_id>/scoring/criteria", methods=["POST"])
@login_required
@require_capture_manager
def add_scoring_criterion(plan_id: int):
    """Add a new scoring criterion."""
    plan = db.get_capture_plan(plan_id)
    if not plan:
        abort(404)
    
    name = request.form.get("name", "").strip()
    if not name:
        flash("Criterion name is required.", "warning")
        return redirect(url_for("scoring_dashboard", plan_id=plan_id))
    
    criterion_id = db.create_scoring_criteria(
        capture_plan_id=plan_id,
        name=name,
        description=request.form.get("description", "").strip() or None,
        weight=float(request.form.get("weight", 1.0)),
        max_score=float(request.form.get("max_score", 10.0)),
        guidance=request.form.get("guidance", "").strip() or None,
        created_by_user_id=current_user.id
    )
    
    if criterion_id:
        db.write_audit_log("SCORING_CRITERIA_CREATED",
                          username=current_user.username,
                          user_id=current_user.id,
                          detail=f"Added criterion '{name}' to capture plan",
                          ip_address=request.remote_addr)
        flash(f"Criterion '{name}' created.", "success")
    else:
        flash("Failed to create criterion.", "danger")
    
    return redirect(url_for("scoring_dashboard", plan_id=plan_id))


@app.route("/capture-plans/<int:plan_id>/scoring/criteria/<int:crit_id>/delete", methods=["POST"])
@login_required
@require_capture_manager
def delete_scoring_criterion(plan_id: int, crit_id: int):
    """Delete a scoring criterion."""
    plan = db.get_capture_plan(plan_id)
    if not plan:
        abort(404)
    
    if db.delete_scoring_criteria(crit_id):
        db.write_audit_log("SCORING_CRITERIA_DELETED",
                          username=current_user.username,
                          user_id=current_user.id,
                          detail="Deleted scoring criterion",
                          ip_address=request.remote_addr)
        flash("Criterion deleted.", "success")
    else:
        flash("Failed to delete criterion.", "danger")
    
    return redirect(url_for("scoring_dashboard", plan_id=plan_id))


@app.route("/projects/<int:project_id>/scoring", methods=["GET"])
@login_required
def score_proposal_page(project_id: int):
    """View scoring interface for a proposal."""
    project = db.get_project(project_id)
    if not project:
        flash("Project not found.", "warning")
        return redirect(url_for("projects"))
    
    # Get capture plan if linked
    capture_plan_id = project.get('capture_plan_id')
    if not capture_plan_id:
        flash("Project not linked to a capture plan.", "warning")
        return redirect(url_for("project_detail", project_id=project_id))
    
    criteria = db.get_scoring_criteria(capture_plan_id)
    scores_data = db.get_proposal_scores(project_id)
    final_score = db.calculate_final_score(project_id)
    ranking = db.get_proposal_ranking(project_id)
    
    return render_template("score_proposal.html",
                          project=project,
                          criteria=criteria,
                          scores=scores_data,
                          final_score=final_score,
                          ranking=ranking)


@app.route("/projects/<int:project_id>/scoring/score", methods=["POST"])
@login_required
def submit_proposal_score(project_id: int):
    """Submit a score for a proposal."""
    project = db.get_project(project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    criterion_id = request.form.get("criterion_id", type=int)
    score_value = request.form.get("score_value", type=float)
    comments = request.form.get("comments", "").strip() or None
    
    if criterion_id is None or score_value is None:
        flash("Criterion and score value are required.", "warning")
        return redirect(url_for("score_proposal_page", project_id=project_id))
    
    if db.score_proposal(project_id, criterion_id, score_value, comments, current_user.id):
        # Recalculate rankings
        if project.get('capture_plan_id'):
            db.recalculate_rankings(project['capture_plan_id'])
        
        db.write_audit_log("PROPOSAL_SCORED",
                          username=current_user.username,
                          user_id=current_user.id,
                          detail=f"Scored proposal '{project['name']}'",
                          ip_address=request.remote_addr)
        
        flash("Score saved.", "success")
    else:
        flash("Failed to save score.", "danger")
    
    return redirect(url_for("score_proposal_page", project_id=project_id))


@app.route("/capture-plans/<int:plan_id>/rankings", methods=["GET"])
@login_required
@require_capture_manager
def view_rankings(plan_id: int):
    """View proposal rankings for a capture plan."""
    plan = db.get_capture_plan(plan_id)
    if not plan:
        flash("Capture plan not found.", "warning")
        return redirect(url_for("capture_dashboard"))
    
    sort_by = request.args.get("sort_by", "rank")
    rankings = db.get_capture_plan_rankings(plan_id, sort_by)
    progress = db.get_scoring_progress(plan_id)
    
    return render_template("proposal_rankings.html",
                          plan=plan,
                          rankings=rankings,
                          progress=progress,
                          sort_by=sort_by)


@app.route("/api/projects/<int:project_id>/scores", methods=["GET"])
@login_required
def get_proposal_scores_api(project_id: int):
    """Get all scores for a proposal (JSON API)."""
    project = db.get_project(project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    scores_data = db.get_proposal_scores(project_id)
    return jsonify(scores_data)


@app.route("/api/projects/<int:project_id>/final-score", methods=["GET"])
@login_required
def get_final_score_api(project_id: int):
    """Get calculated final score (JSON API)."""
    project = db.get_project(project_id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    
    final_score = db.calculate_final_score(project_id)
    ranking = db.get_proposal_ranking(project_id)
    
    return jsonify({
        'final_score': final_score,
        'ranking': dict(ranking) if ranking else None
    })


@app.route("/api/capture-plans/<int:plan_id>/rankings", methods=["GET"])
@login_required
def get_rankings_api(plan_id: int):
    """Get rankings as JSON."""
    plan = db.get_capture_plan(plan_id)
    if not plan:
        return jsonify({'error': 'Capture plan not found'}), 404
    
    sort_by = request.args.get("sort_by", "rank")
    rankings = db.get_capture_plan_rankings(plan_id, sort_by)
    
    return jsonify({
        'plan_id': plan_id,
        'rankings': [dict(r) for r in rankings],
        'count': len(rankings)
    })


# ── Task Management ────────────────────────────────────────────────────────────

@app.route("/tasks")
@login_required
def tasks():
    """Task management page."""
    db.expire_stale_tasks()  # auto-transition expired tasks

    view    = request.args.get("view", "mine")       # mine | assigned | all | project
    proj_id = request.args.get("project_id", type=int)
    status  = request.args.get("status", "")

    if view == "assigned":
        task_list = db.get_tasks(assigned_to_id=current_user.id,
                                  status=status or None)
    elif view == "all":
        task_list = db.get_tasks(status=status or None)
    elif view == "project" and proj_id:
        task_list = db.get_tasks(project_id=proj_id, status=status or None)
    else:  # mine — created by or assigned to me
        task_list = db.get_tasks(user_id=current_user.id, status=status or None)

    projects = db.get_projects(limit=200)
    users    = db.get_all_users()
    counts   = db.get_task_counts_for_user(current_user.id)

    return render_template("tasks.html",
                           task_list=task_list,
                           projects=projects,
                           users=users,
                           counts=counts,
                           view=view,
                           sel_project_id=proj_id,
                           sel_status=status)


@app.route("/tasks/create", methods=["POST"])
@login_required
def create_task():
    data = {
        "title":          request.form.get("title", "").strip(),
        "description":    request.form.get("description", "").strip(),
        "project_id":     request.form.get("project_id") or None,
        "deliverable":    request.form.get("deliverable", "").strip(),
        "created_by_id":  current_user.id,
        "assigned_to_id": request.form.get("assigned_to_id") or None,
        "start_date":     request.form.get("start_date") or None,
        "end_date":       request.form.get("end_date") or None,
        "expire_date":    request.form.get("expire_date") or None,
        "status":         "active",
        "priority":       request.form.get("priority", "normal"),
    }
    if not data["title"]:
        flash("Task title is required.", "warning")
        return redirect(url_for("tasks"))

    task_id = db.create_task(data)

    # Notify assigned user if different from creator
    if data["assigned_to_id"] and int(data["assigned_to_id"]) != current_user.id:
        db.create_notification(
            user_id=int(data["assigned_to_id"]),
            ntype="task_assigned",
            project_id=data["project_id"],
            actor_user_id=current_user.id,
            message=f"{current_user.username} assigned you a task: \"{data['title']}\""
        )

    flash("Task created.", "success")
    return redirect(url_for("tasks"))


@app.route("/tasks/<int:task_id>/edit", methods=["POST"])
@login_required
def edit_task(task_id: int):
    task = db.get_task(task_id)
    if not task:
        abort(404)
    # Only creator or admin can edit
    if task["created_by_id"] != current_user.id and not current_user.is_admin:
        flash("You can only edit tasks you created.", "danger")
        return redirect(url_for("tasks"))

    old_assignee = task.get("assigned_to_id")
    new_assignee = request.form.get("assigned_to_id") or None
    if new_assignee:
        new_assignee = int(new_assignee)

    data = {
        "title":          request.form.get("title", "").strip(),
        "description":    request.form.get("description", "").strip(),
        "project_id":     request.form.get("project_id") or None,
        "deliverable":    request.form.get("deliverable", "").strip(),
        "assigned_to_id": new_assignee,
        "start_date":     request.form.get("start_date") or None,
        "end_date":       request.form.get("end_date") or None,
        "expire_date":    request.form.get("expire_date") or None,
        "status":         request.form.get("status", task["status"]),
        "priority":       request.form.get("priority", "normal"),
    }
    db.update_task(task_id, data)

    # Notify if assignee changed
    if new_assignee and new_assignee != old_assignee and new_assignee != current_user.id:
        db.create_notification(
            user_id=new_assignee,
            ntype="task_assigned",
            project_id=data["project_id"],
            actor_user_id=current_user.id,
            message=f"{current_user.username} assigned you a task: \"{data['title']}\""
        )

    flash("Task updated.", "success")
    return redirect(url_for("tasks"))


@app.route("/tasks/<int:task_id>/status", methods=["POST"])
@login_required
def update_task_status(task_id: int):
    task = db.get_task(task_id)
    if not task:
        abort(404)
    new_status = request.form.get("status", "active")
    db.update_task(task_id, {**task, "status": new_status})
    return redirect(request.referrer or url_for("tasks"))


@app.route("/tasks/<int:task_id>/delete", methods=["POST"])
@login_required
def delete_task(task_id: int):
    task = db.get_task(task_id)
    if not task:
        abort(404)
    if task["created_by_id"] != current_user.id and not current_user.is_admin:
        flash("You can only delete tasks you created.", "danger")
        return redirect(url_for("tasks"))
    db.delete_task(task_id)
    flash("Task deleted.", "success")
    return redirect(url_for("tasks"))


@app.route("/projects/<int:project_id>/checklist/<int:item_id>/schedule", methods=["POST"])
@login_required
def update_checklist_schedule(project_id: int, item_id: int):
    """Update scheduling fields (assignee, dates, hours) for a checklist item."""
    db.update_checklist_schedule(
        item_id=item_id,
        assigned_to_id=request.form.get("assigned_to_id") or None,
        start_date=request.form.get("start_date") or None,
        end_date=request.form.get("end_date") or None,
        estimated_hours=request.form.get("estimated_hours") or 0,
        actual_hours=request.form.get("actual_hours") or 0,
    )
    return redirect(url_for("project_detail", project_id=project_id))
